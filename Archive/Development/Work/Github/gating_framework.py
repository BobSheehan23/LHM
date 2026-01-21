import pandas as pd
import numpy as np
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from datetime import datetime
import matplotlib.pyplot as plt
import io

# --- Configuration & Setup ---

# Constants for styling the Word document
MAIN_COLOR = RGBColor(0x00, 0x6E, 0xB7)
FONT_NAME = 'Calibri'
DATA_FILE_PATH = '/Users/bob/Desktop/Github/Daily_Data_Dispatch_2025-06-17_bsheehan.csv' # Use a descriptive variable
REPORT_DATE = datetime.today()

# --- Helper Functions for Document Generation ---

def add_colored_heading(doc, text, level):
    """Adds a heading with a specific color and font to the document."""
    run = doc.add_heading(level=level).add_run(text)
    run.font.color.rgb = MAIN_COLOR
    run.font.name = FONT_NAME

def add_paragraph(doc, text, size=11, is_bullet=False):
    """Adds a styled paragraph to the document."""
    style = 'List Bullet' if is_bullet else 'Normal'
    p = doc.add_paragraph(text, style=style)
    p.style.font.name = FONT_NAME
    p.style.font.size = Pt(size)
    return p

# --- Core Analysis Functions ---

def load_and_prepare_data(file_path):
    """
    Loads data from the specified CSV and engineers necessary features.
    Handles potential errors during file loading and data processing.
    """
    try:
        df = pd.read_csv(file_path)
    except FileNotFoundError:
        print(f"Error: The file '{file_path}' was not found.")
        return None

    # Rename columns for easier access
    df.rename(columns={
        'Utilization (%)': 'Utilization_Pct',
        'On Loan Value (USD)': 'OnLoanValueUSD',
        'Short Interest Indicator': 'ShortInterestPct',
        'Total Lendable Value (USD)': 'LendableValueUSD',
        'Security Price (USD)': 'PriceUSD',
        'Borrower Count': 'BorrowerCount',
        'Lender Count': 'LenderCount',
        'Fee All (BPS)': 'AvgFee',
        'On Loan Quantity': 'OnLoanQty',
        'Days To Cover': 'DaysToCover'
    }, inplace=True, errors='ignore') # ignore errors if columns don't exist

    # --- Feature Engineering with Error Handling ---
    # Calculate Float
    # Avoid division by zero if ShortInterestPct is 0 or NaN
    df['Float'] = np.where(df['ShortInterestPct'] > 0, df['OnLoanQty'] / (df['ShortInterestPct'] / 100), 0)

    # Calculate Lendable Shares
    df['LendableShares'] = np.where(df['PriceUSD'] > 0, df['LendableValueUSD'] / df['PriceUSD'], 0)

    # Calculate Lendable as % of Float
    df['LendablePctFloat'] = np.where(df['Float'] > 0, (df['LendableShares'] / df['Float']) * 100, 0)

    # Calculate Market Cap for richer analysis
    df['MarketCapUSD'] = df['Float'] * df['PriceUSD']

    return df

def apply_gating_framework(df):
    """
    Applies a sequence of gates to the DataFrame, tracking the impact of each gate.
    Returns a dictionary of results and a list of filtered-out securities' characteristics.
    """
    if df is None:
        return {}, [], None

    # Define the gates in the order they should be applied
    gates = {
        '1. Initial Universe': lambda d: d,
        '2. On-Loan Value >= $1M': lambda d: d[d['OnLoanValueUSD'] >= 1_000_000],
        '3. Lendable > 10% of Float': lambda d: d[d['LendablePctFloat'] > 10],
        '4. Utilization >= 50%': lambda d: d[d['Utilization_Pct'] >= 50],
        '5. Days to Cover >= 2': lambda d: d[d['DaysToCover'] >= 2], # Added this robust gate
        '6. Borrower Count >= 3': lambda d: d[d['BorrowerCount'] >= 3],
        '7. Lender Count >= 2': lambda d: d[d['LenderCount'] >= 2]
    }

    gate_counts = {}
    dropped_securities_analysis = []

    df_filtered = df.copy()
    last_count = len(df_filtered)

    for name, gate_func in gates.items():
        df_after_gate = gate_func(df_filtered)
        current_count = len(df_after_gate)
        gate_counts[name] = current_count

        # Analyze the securities that were dropped by this gate
        if current_count < last_count:
            dropped_indices = df_filtered.index.difference(df_after_gate.index)
            dropped_df = df_filtered.loc[dropped_indices]

            analysis = {
                'Gate': name,
                'Dropped Count': last_count - current_count,
                'Median Market Cap (M)': round(dropped_df['MarketCapUSD'].median() / 1e6, 1),
                'Median Utilization (%)': round(dropped_df['Utilization_Pct'].median(), 1),
                'Median Fee (bps)': round(dropped_df['AvgFee'].median(), 1)
            }
            dropped_securities_analysis.append(analysis)

        df_filtered = df_after_gate
        last_count = current_count

    return gate_counts, dropped_securities_analysis, df_filtered


def generate_impact_chart(gate_counts):
    """Generates a bar chart visualizing the gate attrition and returns it as a bytes object."""
    labels = [label.split('. ')[1] for label in gate_counts.keys()]
    counts = list(gate_counts.values())

    fig, ax = plt.subplots(figsize=(10, 6))
    bars = ax.bar(labels, counts, color='#006EB7', alpha=0.8)

    ax.set_ylabel('Number of Securities')
    ax.set_title('Gating Framework Attrition', fontsize=16, pad=20)
    ax.spines['top'].set_visible(False)
    ax.spines['right'].set_visible(False)
    ax.tick_params(axis='x', rotation=45, labelsize=10)
    ax.grid(axis='y', linestyle='--', alpha=0.6)

    # Add count labels on top of bars
    for bar in bars:
        yval = bar.get_height()
        ax.text(bar.get_x() + bar.get_width()/2.0, yval + counts[0]*0.01, f'{yval:,}', ha='center', va='bottom')

    plt.tight_layout()

    # Save plot to a memory buffer
    img_buffer = io.BytesIO()
    plt.savefig(img_buffer, format='png', dpi=300)
    img_buffer.seek(0)
    plt.close(fig) # Close the figure to free memory

    return img_buffer


def create_report_document(gate_counts, dropped_analysis, final_list_df):
    """
    Generates the complete Word document report.
    """
    doc = Document()

    # --- Header ---
    add_colored_heading(doc, 'EquiLend Data & Analytics — Gating Framework Analysis', 0)
    add_paragraph(doc, 'Date: ' + REPORT_DATE.strftime('%d %b %Y'))

    # --- Methodology ---
    add_colored_heading(doc, '1. Sequential Gating Methodology', 1)
    add_paragraph(doc, 'The following sequential filters are applied to identify securities with the highest potential for short-squeeze risk. Each gate is designed to eliminate common "false positives" and refine the list to only the most compelling candidates.')

    # Gate definitions
    gate_rationales = [
        "On-Loan Value ≥ $1M: Ensures the security has a meaningful level of short interest.",
        "Lendable Inventory > 10% of Float: Confirms the security is part of the liquid lendable market, avoiding signals driven by artificially small supply.",
        "Utilization ≥ 50%: Focuses on names where lendable supply is significantly constrained.",
        "Days to Cover ≥ 2: A critical measure indicating it would take multiple days for shorts to exit positions, which can amplify a squeeze.",
        "Borrower Count ≥ 3: Validates that short demand is broad-based and not driven by a single entity.",
        "Lender Count ≥ 2: Ensures the supply side is not concentrated, reducing the risk of a single lender's actions creating a misleading signal."
    ]
    for rationale in gate_rationales:
        add_paragraph(doc, rationale, is_bullet=True)

    # --- Impact Analysis ---
    add_colored_heading(doc, '2. Impact on Universe Size', 1)

    # Add summary text
    initial_count = list(gate_counts.values())[0]
    final_count = list(gate_counts.values())[-1]
    reduction_pct = round(100 * (1 - final_count / initial_count), 1) if initial_count > 0 else 0
    add_paragraph(doc, f'The framework reduced an initial universe of {initial_count:,} securities to a final qualified list of {final_count:,} (-{reduction_pct}%). The chart below illustrates the attrition at each stage.')

    # Add the attrition chart
    chart_buffer = generate_impact_chart(gate_counts)
    doc.add_picture(chart_buffer, width=Inches(6.0))

    # --- Analysis of Dropped Securities ---
    add_colored_heading(doc, '3. Profile of Filtered Securities', 1)
    add_paragraph(doc, 'Analyzing the characteristics of securities dropped at each gate provides insight into the function of each filter.')

    if dropped_analysis:
        table = doc.add_table(rows=1, cols=5)
        table.style = 'Table Grid'
        hdr_cells = table.rows[0].cells
        hdr_cells[0].text = 'Filtering Gate'
        hdr_cells[1].text = 'Securities Dropped'
        hdr_cells[2].text = 'Median Mkt Cap ($M)'
        hdr_cells[3].text = 'Median Utilization (%)'
        hdr_cells[4].text = 'Median Fee (bps)'

        for item in dropped_analysis:
            row_cells = table.add_row().cells
            row_cells[0].text = item['Gate'].split('. ')[1]
            row_cells[1].text = f"{item['Dropped Count']:,}"
            row_cells[2].text = str(item['Median Market Cap (M)'])
            row_cells[3].text = str(item['Median Utilization (%)'])
            row_cells[4].text = str(item['Median Fee (bps)'])

    # --- Final Candidate List ---
    add_colored_heading(doc, '4. Final Candidate List', 1)
    add_paragraph(doc, f"The following {len(final_list_df)} securities passed all gates.")
    if not final_list_df.empty:
        # Add a table of the top candidates
        display_cols = ['TICKER', 'OnLoanValueUSD', 'Utilization_Pct', 'DaysToCover', 'AvgFee', 'BorrowerCount']
        final_list_display = final_list_df[display_cols].sort_values(by='DaysToCover', ascending=False).head(10)

        # Format for display
        final_list_display['OnLoanValueUSD'] = final_list_display['OnLoanValueUSD'].apply(lambda x: f"${x/1e6:,.1f}M")
        final_list_display['Utilization_Pct'] = final_list_display['Utilization_Pct'].apply(lambda x: f"{x:.1f}%")
        final_list_display['DaysToCover'] = final_list_display['DaysToCover'].apply(lambda x: f"{x:.2f}")

        table = doc.add_table(rows=1, cols=len(display_cols))
        table.style = 'Table Grid'
        hdr_cells = table.rows[0].cells
        for i, col_name in enumerate(final_list_display.columns):
            hdr_cells[i].text = col_name

        for index, row in final_list_display.iterrows():
            row_cells = table.add_row().cells
            for i, val in enumerate(row):
                row_cells[i].text = str(val)


    # --- Save the document ---
    file_name = f'EquiLend_Gating_Analysis_{REPORT_DATE.strftime("%Y-%m-%d")}.docx'
    doc.save(file_name)
    print(f"Successfully generated report: '{file_name}'")


# --- Main Execution ---
if __name__ == "__main__":
    # 1. Load and prepare the data
    raw_df = load_and_prepare_data(DATA_FILE_PATH)

    if raw_df is not None:
        # 2. Apply the gating framework
        gate_results, dropped_stats, final_df = apply_gating_framework(raw_df)

        # 3. Generate the final Word document report
        create_report_document(gate_results, dropped_stats, final_df)
