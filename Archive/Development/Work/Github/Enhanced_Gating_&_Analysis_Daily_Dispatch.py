# %% [markdown]
# # EquiLend Data & Analytics | Daily Dispatch Template (v2)
# 
# ## Market Notes
# - Average Fee (BPS): [Enter Value]
# - Average Fee Change (%): [Enter Value]
# - Average Utilization (%): [Enter Value]
# - Average Utilization Change (%): [Enter Value]
# 
# ## Sectors in Focus
# - [Enter Sector 1]
# - [Enter Sector 2]
# - [Enter Sector 3]
# 
# ## Top 5 Headlines
# 1. [Enter Headline 1]
# 2. [Enter Headline 2]
# 3. [Enter Headline 3]
# 4. [Enter Headline 4]
# 5. [Enter Headline 5]
# 
# ## Data Tables
# - [Specials/Hard-to-Borrow Table Placeholder]
# - [Global Short Squeeze Score Table Placeholder]
# 
# ## Key Takeaways
# - [Enter Takeaway 1]
# - [Enter Takeaway 2]
# - [Enter Takeaway 3]

# %% [markdown]
# # Gating Framework: Thought Process, Research, and Overview
# 
# ## Overview
# The gating framework is a systematic, data-driven process designed to identify securities with the highest potential for short-squeeze risk. By applying a series of sequential filters ("gates"), the framework narrows down a broad universe of securities to a focused list of candidates that meet strict criteria for liquidity, demand, and risk factors.
# 
# ## Gate Sequence and Rationale
# 1. **Initial Universe**: All securities in the dataset.
# 2. **On-Loan Value ≥ $1M**: Ensures the security has a meaningful level of short interest, filtering out illiquid or inactive names.
# 3. **Borrower Count ≥ 3**: Validates that short demand is broad-based and not driven by a single entity, reducing idiosyncratic risk.
# 4. **Lender Count ≥ 2**: Ensures the supply side is not concentrated, reducing the risk of a single lender's actions creating a misleading signal.
# 5. **Lendable Inventory > 10% of Float**: Confirms the security is part of the liquid lendable market, avoiding signals driven by artificially small supply.
# 6. **Utilization ≥ 50%**: Focuses on names where lendable supply is significantly constrained, increasing squeeze potential.
# 7. **Days to Cover ≥ 2**: Indicates it would take multiple days for shorts to exit positions, which can amplify a squeeze.
# 
# ## Float Calculation
# Float is a critical metric in the analysis, calculated as:
# 
#     Float = On Loan Quantity / (Short Interest Indicator / 100)
# 
# This approach ensures that the float reflects the true available supply in the market, accounting for both borrow activity and reported short interest.
# 
# ## Research & Thought Process
# - **Liquidity and Market Depth**: The initial gates focus on ensuring that only liquid, actively traded securities are considered, as these are more likely to experience meaningful price movements.
# - **Breadth of Demand and Supply**: By requiring a minimum number of borrowers and lenders, the framework avoids names where activity is dominated by a single participant, which could skew results or create false signals.
# - **Supply Constraints**: High utilization and low lendable inventory relative to float are classic indicators of potential squeezes, as they signal that shorts may struggle to cover positions if demand spikes or supply contracts.
# - **Exit Difficulty**: Days to cover is a direct measure of how hard it would be for shorts to unwind positions, with higher values indicating greater risk of a squeeze.
# 
# ## Summary
# This gating process is the result of research into market microstructure, historical short-squeeze events, and best practices in securities lending analytics. The framework is designed to be robust, transparent, and adaptable to evolving market conditions.

# %% [markdown]
# <a href="https://colab.research.google.com/github/BobSheehan23/EquiLend/blob/main/Enhanced_Gating_%26_Analysis_Framework.ipynb" target="_parent"><img src="https://colab.research.google.com/assets/colab-badge.svg" alt="Open In Colab"/></a>

# %%
import pandas as pd
import numpy as np
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from datetime import datetime
import matplotlib.pyplot as plt
import io

# === Configuration ===
MAIN_COLOR = RGBColor(0x00, 0x6E, 0xB7)
FONT_NAME = 'Calibri'
DATA_FILE_PATH = 'Daily_Data_Dispatch_2025-06-16_bsheehan_adhoc.csv'
REPORT_DATE = datetime.today()

# === Word Document Helpers ===
def add_colored_heading(doc, text, level):
    run = doc.add_heading(level=level).add_run(text)
    run.font.color.rgb = MAIN_COLOR
    run.font.name = FONT_NAME

def add_paragraph(doc, text, size=11, is_bullet=False):
    style = 'List Bullet' if is_bullet else 'Normal'
    p = doc.add_paragraph(text, style=style)
    p.style.font.name = FONT_NAME
    p.style.font.size = Pt(size)
    return p

# === Data Preparation ===
def load_and_prepare_data(file_path):
    try:
        df = pd.read_csv(file_path)
    except FileNotFoundError:
        print(f"Error: The file '{file_path}' was not found.")
        return None
    df.rename(columns={
        'Utilization (%)': 'Utilization_Pct',
        'On Loan Value (USD)': 'OnLoanValueUSD',
        'Short Interest Indicator': 'ShortInterestPct',
        'Total Lendable Value (USD)': 'LendableValueUSD',
        'Security Price (USD)': 'PriceUSD',
        'Borrower Count': 'BorrowerCount',
        'Lender Count': 'LenderCount',
        'Average Fee': 'AvgFee',
        'On Loan Quantity': 'OnLoanQty',
        'Composite 20-Day ADV': 'ADV20Day'
    }, inplace=True, errors='ignore')
    df['Float'] = np.where(df['ShortInterestPct'] > 0, df['OnLoanQty'] / (df['ShortInterestPct'] / 100), 0)
    df['LendableShares'] = np.where(df['PriceUSD'] > 0, df['LendableValueUSD'] / df['PriceUSD'], 0)
    df['LendablePctFloat'] = np.where(df['Float'] > 0, (df['LendableShares'] / df['Float']) * 100, 0)
    df['DaysToCover'] = np.where(df['ADV20Day'] > 0, df['OnLoanQty'] / df['ADV20Day'], 0)
    df['MarketCapUSD'] = df['Float'] * df['PriceUSD']
    return df

# === Gating Framework ===
def apply_gating_framework(df):
    if df is None:
        return {}, [], pd.DataFrame()
    gates = {
        '1. Initial Universe': lambda d: d,
        '2. On-Loan Value >= $1M': lambda d: d[d['OnLoanValueUSD'] >= 1_000_000],
        '3. Lendable > 10% of Float': lambda d: d[d['LendablePctFloat'] > 10],
        '4. Utilization >= 50%': lambda d: d[d['Utilization_Pct'] >= 50],
        '5. Days to Cover >= 2': lambda d: d[d['DaysToCover'] >= 2],
        '6. Borrower Count >= 3': lambda d: d[d['BorrowerCount'] >= 3],
        '7. Lender Count >= 2': lambda d: d[d['LenderCount'] >= 2]
    }
    gate_counts = {}
    dropped_securities_analysis = []
    df_filtered = df.copy()
    last_count = len(df_filtered)
    for name, gate_func in gates.items():
        df_after_gate = gate_func(df_filtered)
        current_count = len(df_after_gate)
        gate_counts[name] = current_count
        if current_count < last_count:
            dropped_indices = df_filtered.index.difference(df_after_gate.index)
            dropped_df = df_filtered.loc[dropped_indices]
            analysis = {
                'Gate': name,
                'Dropped Count': last_count - current_count,
                'Median Market Cap (M)': round(dropped_df['MarketCapUSD'].median() / 1e6, 1),
                'Median Utilization (%)': round(dropped_df['Utilization_Pct'].median(), 1),
                'Median Fee (bps)': round(dropped_df['AvgFee'].median(), 1)
            }
            dropped_securities_analysis.append(analysis)
        df_filtered = df_after_gate
        last_count = current_count
    return gate_counts, dropped_securities_analysis, df_filtered

# === Visualization ===
def generate_impact_chart(gate_counts):
    labels = [label.split('. ')[1] for label in gate_counts.keys()]
    counts = list(gate_counts.values())
    fig, ax = plt.subplots(figsize=(10, 6))
    bars = ax.bar(labels, counts, color='#006EB7', alpha=0.8)
    ax.set_ylabel('Number of Securities')
    ax.set_title('Gating Framework Attrition', fontsize=16, pad=20)
    ax.spines['top'].set_visible(False)
    ax.spines['right'].set_visible(False)
    ax.tick_params(axis='x', rotation=45, labelsize=10)
    ax.grid(axis='y', linestyle='--', alpha=0.6)
    for bar in bars:
        yval = bar.get_height()
        ax.text(bar.get_x() + bar.get_width()/2.0, yval + counts[0]*0.01, f'{yval:,}', ha='center', va='bottom')
    plt.tight_layout()
    img_buffer = io.BytesIO()
    plt.savefig(img_buffer, format='png', dpi=300)
    img_buffer.seek(0)
    plt.close(fig)
    return img_buffer

# === Report Generation ===
def create_report_document(gate_counts, dropped_analysis, final_list_df):
    doc = Document()
    add_colored_heading(doc, 'EquiLend Data & Analytics — Gating Framework Analysis', 0)
    add_paragraph(doc, 'Date: ' + REPORT_DATE.strftime('%d %b %Y'))
    add_colored_heading(doc, '1. Sequential Gating Methodology', 1)
    add_paragraph(doc, 'The following sequential filters are applied to identify securities with the highest potential for short-squeeze risk. Each gate is designed to eliminate common "false positives" and refine the list to only the most compelling candidates.')
    gate_rationales = [
        "On-Loan Value ≥ $1M: Ensures the security has a meaningful level of short interest.",
        "Lendable Inventory > 10% of Float: Confirms the security is part of the liquid lendable market, avoiding signals driven by artificially small supply.",
        "Utilization ≥ 50%: Focuses on names where lendable supply is significantly constrained.",
        "Days to Cover ≥ 2: A critical measure indicating it would take multiple days for shorts to exit positions, which can amplify a squeeze.",
        "Borrower Count ≥ 3: Validates that short demand is broad-based and not driven by a single entity.",
        "Lender Count ≥ 2: Ensures the supply side is not concentrated, reducing the risk of a single lender's actions creating a misleading signal."
    ]
    for rationale in gate_rationales:
        add_paragraph(doc, rationale, is_bullet=True)
    add_colored_heading(doc, '2. Impact on Universe Size', 1)
    initial_count = list(gate_counts.values())[0]
    final_count = list(gate_counts.values())[-1]
    reduction_pct = round(100 * (1 - final_count / initial_count), 1) if initial_count > 0 else 0
    add_paragraph(doc, f'The framework reduced an initial universe of {initial_count:,} securities to a final qualified list of {final_count:,} (-{reduction_pct}%). The chart below illustrates the attrition at each stage.')
    chart_buffer = generate_impact_chart(gate_counts)
    doc.add_picture(chart_buffer, width=Inches(6.0))
    add_colored_heading(doc, '3. Profile of Filtered Securities', 1)
    add_paragraph(doc, 'Analyzing the characteristics of securities dropped at each gate provides insight into the function of each filter.')
    if dropped_analysis:
        table = doc.add_table(rows=1, cols=5)
        table.style = 'Table Grid'
        hdr_cells = table.rows[0].cells
        hdr_cells[0].text = 'Filtering Gate'
        hdr_cells[1].text = 'Securities Dropped'
        hdr_cells[2].text = 'Median Mkt Cap ($M)'
        hdr_cells[3].text = 'Median Utilization (%)'
        hdr_cells[4].text = 'Median Fee (bps)'
        for item in dropped_analysis:
            row_cells = table.add_row().cells
            row_cells[0].text = item['Gate'].split('. ')[1]
            row_cells[1].text = f"{item['Dropped Count']:,}"
            row_cells[2].text = str(item['Median Market Cap (M)'])
            row_cells[3].text = str(item['Median Utilization (%)'])
            row_cells[4].text = str(item['Median Fee (bps)'])
    add_colored_heading(doc, '4. Final Candidate List', 1)
    add_paragraph(doc, f"The following {len(final_list_df)} securities passed all gates.")
    if not final_list_df.empty:
        display_cols = ['Ticker', 'OnLoanValueUSD', 'Utilization_Pct', 'DaysToCover', 'AvgFee', 'BorrowerCount']
        final_list_display = final_list_df[display_cols].sort_values(by='DaysToCover', ascending=False).head(10)
        final_list_display['OnLoanValueUSD'] = final_list_display['OnLoanValueUSD'].apply(lambda x: f"${x/1e6:,.1f}M")
        final_list_display['Utilization_Pct'] = final_list_display['Utilization_Pct'].apply(lambda x: f"{x:.1f}%")
        final_list_display['DaysToCover'] = final_list_display['DaysToCover'].apply(lambda x: f"{x:.2f}")
        table = doc.add_table(rows=1, cols=len(display_cols))
        table.style = 'Table Grid'
        hdr_cells = table.rows[0].cells
        for i, col_name in enumerate(final_list_display.columns):
            hdr_cells[i].text = col_name
        for index, row in final_list_display.iterrows():
            row_cells = table.add_row().cells
            for i, val in enumerate(row):
                row_cells[i].text = str(val)
    file_name = f'EquiLend_Gating_Analysis_{REPORT_DATE.strftime("%Y-%m-%d")}.docx'
    doc.save(file_name)
    print(f"Successfully generated report: '{file_name}'")

# === Main Execution ===
if __name__ == "__main__":
    raw_df = load_and_prepare_data(DATA_FILE_PATH)
    if raw_df is not None:
        gate_results, dropped_stats, final_df = apply_gating_framework(raw_df)
        create_report_document(gate_results, dropped_stats, final_df)


