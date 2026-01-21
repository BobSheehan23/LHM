"""
The Horizon Report - January 2026
Final Institutional-Quality Report Generator

Lighthouse Macro
MACRO, ILLUMINATED.

This script generates a complete DOCX report with:
- 36 original Horizon charts
- 7 new Priority 1 charts
- Full analytical content
- Institutional-quality formatting
"""

from docx import Document
from docx.shared import Inches, Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from pathlib import Path
from datetime import datetime
import os

# Paths
CHART_DIR = Path("/Users/bob/LHM/Images/horizon charts")
PRIORITY1_DIR = Path("/Users/bob/LHM/data/lighthouse_data/charts/priority1")
OUTPUT_DIR = Path("/Users/bob/LHM/deliverables/horizon_report_final")
OUTPUT_DIR.mkdir(parents=True, exist_ok=True)

# Lighthouse Colors
OCEAN_BLUE = RGBColor(0x00, 0x89, 0xD1)
DUSK_ORANGE = RGBColor(0xFF, 0x67, 0x23)
NEUTRAL_GRAY = RGBColor(0x80, 0x80, 0x80)
BLACK = RGBColor(0x00, 0x00, 0x00)


def set_cell_shading(cell, color_hex):
    """Set cell background color."""
    shading_elm = OxmlElement('w:shd')
    shading_elm.set(qn('w:fill'), color_hex)
    cell._tc.get_or_add_tcPr().append(shading_elm)


def add_horizontal_line(doc):
    """Add a horizontal line."""
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after = Pt(12)
    run = p.add_run("─" * 80)
    run.font.color.rgb = NEUTRAL_GRAY
    run.font.size = Pt(8)


def add_page_break(doc):
    """Add page break."""
    doc.add_page_break()


def add_chart(doc, chart_path, caption=None, width=6.5):
    """Add a chart image with optional caption."""
    if chart_path.exists():
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run()
        run.add_picture(str(chart_path), width=Inches(width))

        if caption:
            cap = doc.add_paragraph(caption)
            cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
            cap.runs[0].font.size = Pt(9)
            cap.runs[0].font.italic = True
            cap.runs[0].font.color.rgb = NEUTRAL_GRAY
        return True
    return False


def create_document():
    """Create the main document with styles."""
    doc = Document()

    # Set default margins
    for section in doc.sections:
        section.top_margin = Inches(0.75)
        section.bottom_margin = Inches(0.75)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)

    # Define styles
    styles = doc.styles

    # Title style
    title_style = styles.add_style('LHM Title', WD_STYLE_TYPE.PARAGRAPH)
    title_style.font.name = 'Arial'
    title_style.font.size = Pt(28)
    title_style.font.bold = True
    title_style.font.color.rgb = OCEAN_BLUE
    title_style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_style.paragraph_format.space_after = Pt(6)

    # Subtitle style
    sub_style = styles.add_style('LHM Subtitle', WD_STYLE_TYPE.PARAGRAPH)
    sub_style.font.name = 'Arial'
    sub_style.font.size = Pt(14)
    sub_style.font.color.rgb = NEUTRAL_GRAY
    sub_style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub_style.paragraph_format.space_after = Pt(24)

    # Heading 1 (Part titles)
    h1_style = styles.add_style('LHM H1', WD_STYLE_TYPE.PARAGRAPH)
    h1_style.font.name = 'Arial'
    h1_style.font.size = Pt(18)
    h1_style.font.bold = True
    h1_style.font.color.rgb = OCEAN_BLUE
    h1_style.paragraph_format.space_before = Pt(24)
    h1_style.paragraph_format.space_after = Pt(12)

    # Heading 2 (Section titles)
    h2_style = styles.add_style('LHM H2', WD_STYLE_TYPE.PARAGRAPH)
    h2_style.font.name = 'Arial'
    h2_style.font.size = Pt(14)
    h2_style.font.bold = True
    h2_style.font.color.rgb = BLACK
    h2_style.paragraph_format.space_before = Pt(18)
    h2_style.paragraph_format.space_after = Pt(6)

    # Body text
    body_style = styles.add_style('LHM Body', WD_STYLE_TYPE.PARAGRAPH)
    body_style.font.name = 'Georgia'
    body_style.font.size = Pt(11)
    body_style.font.color.rgb = BLACK
    body_style.paragraph_format.space_after = Pt(8)
    body_style.paragraph_format.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
    body_style.paragraph_format.line_spacing = 1.15

    # Quote style
    quote_style = styles.add_style('LHM Quote', WD_STYLE_TYPE.PARAGRAPH)
    quote_style.font.name = 'Georgia'
    quote_style.font.size = Pt(12)
    quote_style.font.italic = True
    quote_style.font.color.rgb = OCEAN_BLUE
    quote_style.paragraph_format.left_indent = Inches(0.5)
    quote_style.paragraph_format.space_before = Pt(12)
    quote_style.paragraph_format.space_after = Pt(12)

    return doc


def add_cover_page(doc):
    """Add the cover page."""
    # Spacer
    for _ in range(3):
        doc.add_paragraph()

    # Logo/Brand
    p = doc.add_paragraph("LIGHTHOUSE MACRO")
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.runs[0].font.name = 'Arial'
    p.runs[0].font.size = Pt(14)
    p.runs[0].font.bold = True
    p.runs[0].font.color.rgb = OCEAN_BLUE

    p = doc.add_paragraph("MACRO, ILLUMINATED.")
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.runs[0].font.name = 'Arial'
    p.runs[0].font.size = Pt(10)
    p.runs[0].font.color.rgb = NEUTRAL_GRAY

    add_horizontal_line(doc)

    # Main title
    for _ in range(2):
        doc.add_paragraph()

    p = doc.add_paragraph("THE HORIZON REPORT")
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.runs[0].font.name = 'Arial'
    p.runs[0].font.size = Pt(36)
    p.runs[0].font.bold = True
    p.runs[0].font.color.rgb = OCEAN_BLUE

    p = doc.add_paragraph("JANUARY 2026")
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.runs[0].font.name = 'Arial'
    p.runs[0].font.size = Pt(18)
    p.runs[0].font.color.rgb = NEUTRAL_GRAY

    add_horizontal_line(doc)

    # Subtitle
    p = doc.add_paragraph("THE SILENT STOP")
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.runs[0].font.name = 'Arial'
    p.runs[0].font.size = Pt(24)
    p.runs[0].font.bold = True
    p.runs[0].font.color.rgb = BLACK

    p = doc.add_paragraph("Why GDP Is Lying, Liquidity Is Constrained, and Markets Are Fragile")
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.runs[0].font.name = 'Georgia'
    p.runs[0].font.size = Pt(14)
    p.runs[0].font.italic = True
    p.runs[0].font.color.rgb = NEUTRAL_GRAY

    # Author
    for _ in range(4):
        doc.add_paragraph()

    p = doc.add_paragraph("Robert Sheehan, CFA, CMT")
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.runs[0].font.name = 'Arial'
    p.runs[0].font.size = Pt(12)
    p.runs[0].font.bold = True

    p = doc.add_paragraph("Founder & Chief Investment Officer")
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.runs[0].font.name = 'Arial'
    p.runs[0].font.size = Pt(11)
    p.runs[0].font.color.rgb = NEUTRAL_GRAY

    add_page_break(doc)


def add_toc(doc):
    """Add table of contents."""
    p = doc.add_paragraph("TABLE OF CONTENTS", style='LHM H1')
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    add_horizontal_line(doc)

    toc_items = [
        ("EXECUTIVE SUMMARY", "3"),
        ("", ""),
        ("PART I — THE SILENT STOP", "5"),
        ("    I. The GDP–GDI Schism", "6"),
        ("    II. Fiscal Dominance", "8"),
        ("    III. GDI Tells the Truth", "10"),
        ("    IV. Labor: Freezing Not Firing", "12"),
        ("    V. The Consumer: Liquidity to Solvency", "16"),
        ("    VI. K-Shaped Cost of Capital", "20"),
        ("    VII. Defining the Silent Stop", "22"),
        ("", ""),
        ("PART II — MONETARY MECHANICS", "24"),
        ("    I. Liquidity Defined Correctly", "25"),
        ("    II. The Fed's Balance Sheet Identity", "26"),
        ("    III. RRP Exhaustion", "28"),
        ("    IV. Bank Reserves & Adequacy", "30"),
        ("    V. TGA: The Hidden QT Lever", "32"),
        ("    VI. Reserve Management Purchases", "34"),
        ("    VII. Funding Market Stress Signals", "36"),
        ("", ""),
        ("PART III — MARKET TECHNICALS", "40"),
        ("    I. From Plumbing to Price Action", "41"),
        ("    II. Yield Curve Dynamics", "43"),
        ("    III. Treasury Auction Stress", "46"),
        ("    IV. Credit Market Analysis", "50"),
        ("    V. Cross-Asset Signals", "54"),
        ("    VI. Volatility Regime", "56"),
        ("", ""),
        ("CONCLUSION — THE 16-WEEK STRESS WINDOW", "60"),
        ("", ""),
        ("APPENDIX — CHART INDEX", "64"),
    ]

    for item, page in toc_items:
        if item:
            p = doc.add_paragraph()
            p.add_run(item).font.name = 'Arial'
            p.runs[0].font.size = Pt(11)
            if not item.startswith("    "):
                p.runs[0].font.bold = True
            # Add tab and page number
            tab_run = p.add_run("\t" + page)
            tab_run.font.name = 'Arial'
            tab_run.font.size = Pt(11)
            tab_run.font.color.rgb = NEUTRAL_GRAY
        else:
            doc.add_paragraph()

    add_page_break(doc)


def add_executive_summary(doc):
    """Add executive summary section."""
    p = doc.add_paragraph("EXECUTIVE SUMMARY", style='LHM H1')
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    add_horizontal_line(doc)

    # Core thesis
    doc.add_paragraph("THE CORE THESIS: THE SILENT STOP", style='LHM H2')

    doc.add_paragraph(
        "The U.S. economy in early 2026 is not overheating. It is stalling silently. "
        "Headline GDP remains positive (+1.8% YoY), but this figure is propped up entirely "
        "by fiscal injections. The private sector engine—measured by Gross Domestic Income—has "
        "flatlined at 0.0% YoY.",
        style='LHM Body'
    )

    doc.add_paragraph(
        "This is the largest GDP–GDI divergence on record. Historically, such dislocations "
        "resolve via downward GDP revisions—not upward GDI rebounds.",
        style='LHM Body'
    )

    p = doc.add_paragraph(
        '"GDP is the paint job. GDI is the engine. And the engine has seized."',
        style='LHM Quote'
    )

    # Three pillars table
    doc.add_paragraph("THE THREE PILLARS OF FRAGILITY", style='LHM H2')

    # Pillar 1: Macro
    doc.add_paragraph("1. MACRO DYNAMICS: The Big Freeze", style='LHM H2')
    doc.add_paragraph(
        "The labor market appears stable on the surface, but internal dynamism has collapsed:",
        style='LHM Body'
    )

    table = doc.add_table(rows=5, cols=2)
    table.style = 'Table Grid'

    data = [
        ("Indicator", "Current Reading"),
        ("Quits Rate", "1.8% (below 2019 levels)"),
        ("Temp Help Employment (YoY)", "−15% (recession-consistent)"),
        ("Avg Weekly Hours", "34.2 (cycle low)"),
        ("Continuing Claims", "~1.97M (rising trend)"),
    ]

    for i, (col1, col2) in enumerate(data):
        row = table.rows[i]
        row.cells[0].text = col1
        row.cells[1].text = col2
        if i == 0:
            for cell in row.cells:
                cell.paragraphs[0].runs[0].font.bold = True
                set_cell_shading(cell, "0089D1")
                cell.paragraphs[0].runs[0].font.color.rgb = RGBColor(255, 255, 255)

    doc.add_paragraph()
    doc.add_paragraph("The economy isn't laying off. It's locking up.", style='LHM Body')

    # Pillar 2: Monetary
    doc.add_paragraph("2. MONETARY MECHANICS: The Plumbing Break", style='LHM H2')
    doc.add_paragraph(
        "The system's shock absorber has been drained. The Overnight Reverse Repo Facility (RRP), "
        "which absorbed over $2 trillion in excess liquidity in 2022, is now effectively empty:",
        style='LHM Body'
    )

    table2 = doc.add_table(rows=5, cols=2)
    table2.style = 'Table Grid'

    data2 = [
        ("Metric", "Current Reading"),
        ("RRP Balance", "$4.58B (from $2T+ in 2022)"),
        ("Reserve Balances", "~$2.96T (≈$180B above floor)"),
        ("SOFR vs IORB", "SOFR above IORB (stress signal)"),
        ("SRF Usage (Year-End)", "Record levels"),
    ]

    for i, (col1, col2) in enumerate(data2):
        row = table2.rows[i]
        row.cells[0].text = col1
        row.cells[1].text = col2
        if i == 0:
            for cell in row.cells:
                cell.paragraphs[0].runs[0].font.bold = True
                set_cell_shading(cell, "0089D1")
                cell.paragraphs[0].runs[0].font.color.rgb = RGBColor(255, 255, 255)

    doc.add_paragraph()
    doc.add_paragraph(
        "The Fed has quietly launched Reserve Management Purchases (RMP)—Not-QE by another name. "
        "The plumbing has cracked.",
        style='LHM Body'
    )

    # Pillar 3: Markets
    doc.add_paragraph("3. MARKET TECHNICALS: The K-Shaped Stall", style='LHM H2')
    doc.add_paragraph(
        "Monetary policy is now a selective tightening tool. Capital-rich megacaps are insulated "
        "by fixed-rate debt; capital-dependent SMEs are being crushed.",
        style='LHM Body'
    )
    doc.add_paragraph(
        "2026–2027 sees a $480B maturity wall in sub-IG debt. Many marginal firms won't default "
        "from losses—they'll default from inability to refinance.",
        style='LHM Body'
    )

    add_horizontal_line(doc)

    # Dashboard
    doc.add_paragraph("JANUARY 2026 DASHBOARD", style='LHM H2')

    table3 = doc.add_table(rows=9, cols=3)
    table3.style = 'Table Grid'

    dashboard_data = [
        ("Macro Driver", "Status", "Signal"),
        ("GDP", "+1.8% (fiscal-driven)", "CAUTION"),
        ("GDI", "0.0% (private income flat)", "WARNING"),
        ("Labor", "Frozen (quits, hours, temp)", "WARNING"),
        ("Consumer", "Credit stress accelerating", "CAUTION"),
        ("SMEs", "Choked by 9.4%+ borrowing cost", "WARNING"),
        ("Liquidity", "RRP depleted, RMP launched", "WARNING"),
        ("Markets", "Elevated on thin liquidity", "CAUTION"),
        ("Overall", "SILENT STOP REGIME", "ELEVATED RISK"),
    ]

    for i, row_data in enumerate(dashboard_data):
        row = table3.rows[i]
        for j, text in enumerate(row_data):
            row.cells[j].text = text
            if i == 0:
                row.cells[j].paragraphs[0].runs[0].font.bold = True
                set_cell_shading(row.cells[j], "0089D1")
                row.cells[j].paragraphs[0].runs[0].font.color.rgb = RGBColor(255, 255, 255)
            elif j == 2:
                if "WARNING" in text:
                    row.cells[j].paragraphs[0].runs[0].font.color.rgb = RGBColor(255, 103, 35)
                elif "CAUTION" in text:
                    row.cells[j].paragraphs[0].runs[0].font.color.rgb = RGBColor(255, 165, 0)
                elif "ELEVATED" in text:
                    row.cells[j].paragraphs[0].runs[0].font.color.rgb = RGBColor(255, 51, 51)

    add_page_break(doc)


def add_part1(doc):
    """Add Part I - The Silent Stop."""
    # Part header
    p = doc.add_paragraph("PART I", style='LHM H1')
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    p = doc.add_paragraph("THE SILENT STOP")
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.runs[0].font.name = 'Arial'
    p.runs[0].font.size = Pt(24)
    p.runs[0].font.bold = True
    p.runs[0].font.color.rgb = OCEAN_BLUE

    p = doc.add_paragraph("The U.S. Economy Has Not Slowed. It Has Stalled.")
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.runs[0].font.name = 'Georgia'
    p.runs[0].font.size = Pt(12)
    p.runs[0].font.italic = True
    p.runs[0].font.color.rgb = NEUTRAL_GRAY

    add_horizontal_line(doc)

    doc.add_paragraph(
        '"GDP is the paint job. GDI is the engine. And the engine has seized."',
        style='LHM Quote'
    )

    doc.add_paragraph(
        "The defining macro error of the current cycle is not optimism. It is misidentification of motion.",
        style='LHM Body'
    )

    doc.add_paragraph(
        "Headline economic data suggests an economy that is decelerating but intact. GDP growth remains "
        "positive. Employment has softened but not collapsed. Inflation has cooled from its peak. "
        "Financial markets, reading these signals, have extrapolated resilience into permanence. "
        "The dominant narrative is no longer 'soft landing' but 'no landing.'",
        style='LHM Body'
    )

    doc.add_paragraph(
        "That narrative is wrong.",
        style='LHM Body'
    )

    doc.add_paragraph(
        "The U.S. economy in early 2026 is not expanding, slowing, or contracting in the traditional "
        "cyclical sense. It is stalled—held in place by fiscal flow, accounting optics, and residual "
        "liquidity while the private-sector engine that normally generates forward motion has stopped turning.",
        style='LHM Body'
    )

    doc.add_paragraph(
        "This condition is what we define as The Silent Stop.",
        style='LHM Body'
    )

    add_horizontal_line(doc)

    # Section I: GDP-GDI
    doc.add_paragraph("I. THE GDP–GDI SCHISM", style='LHM H2')
    doc.add_paragraph("When the Paint Job Hides the Engine Failure", style='LHM Body')

    doc.add_paragraph(
        "If there is a single macro relationship that should anchor institutional risk assessment in 2026, "
        "it is the divergence between Gross Domestic Product (GDP) and Gross Domestic Income (GDI).",
        style='LHM Body'
    )

    doc.add_paragraph(
        "In theory, these two measures are identical. Every dollar spent is a dollar earned. In practice, "
        "they are constructed from different data sources and diverge at turning points. When they do, "
        "history is unambiguous: GDI leads.",
        style='LHM Body'
    )

    doc.add_paragraph("As of Q4 2025:", style='LHM Body')
    doc.add_paragraph("• Real GDP is growing at approximately +1.8% YoY", style='LHM Body')
    doc.add_paragraph("• Real GDI is effectively 0.0% YoY", style='LHM Body')

    doc.add_paragraph(
        "This is not noise. It is one of the largest sustained divergences on record outside of crisis periods.",
        style='LHM Body'
    )

    # Chart: Real vs Nominal GDP
    add_chart(doc, CHART_DIR / "27.png", "Figure 1.1: Real vs Nominal GDP — Deflator contribution masks real weakness")

    doc.add_paragraph(
        "GDP is an expenditure measure. It captures what is spent, including government outlays financed "
        "by borrowing. GDI is an income measure. It captures what is earned—wages, profits, interest, "
        "and rents accruing to the private sector.",
        style='LHM Body'
    )

    doc.add_paragraph(
        "When GDP exceeds GDI materially and persistently, the implication is not 'hidden strength.' "
        "It is borrowed momentum.",
        style='LHM Body'
    )

    doc.add_paragraph(
        "Historically, these gaps close in one of two ways:",
        style='LHM Body'
    )
    doc.add_paragraph("1. GDI accelerates as private income catches up (rare late-cycle)", style='LHM Body')
    doc.add_paragraph("2. GDP is revised downward as fiscal impulse fades and private weakness asserts itself", style='LHM Body')

    doc.add_paragraph(
        "There is no precedent for sustained GDP growth in the absence of private income growth.",
        style='LHM Body'
    )

    add_page_break(doc)

    # Section II: Fiscal Dominance
    doc.add_paragraph("II. FISCAL DOMINANCE", style='LHM H2')
    doc.add_paragraph("Why GDP Is Lying This Time", style='LHM Body')

    doc.add_paragraph(
        "The reason this divergence is so extreme—and so persistent—is fiscal dominance.",
        style='LHM Body'
    )

    doc.add_paragraph(
        "Government consumption and investment are growing at a pace more consistent with wartime or "
        "crisis conditions than with a mature expansion. Healthcare spending, defense outlays, and "
        "infrastructure disbursements are providing a steady inflow into GDP calculations regardless "
        "of private-sector conditions.",
        style='LHM Body'
    )

    # Chart: Fiscal Dominance Cascade
    add_chart(doc, CHART_DIR / "12.png", "Figure 1.2: The Fiscal Dominance Cascade — Self-reinforcing debt spiral")

    doc.add_paragraph(
        "This creates the illusion of motion. From a national accounting perspective, this is legitimate. "
        "From a macro-cyclical perspective, it is deeply misleading.",
        style='LHM Body'
    )

    doc.add_paragraph(
        "Government spending does not create self-reinforcing economic loops. It does not generate "
        "labor mobility, capital formation, or productivity growth in the same way private investment does. "
        "It sustains activity, but it does not propagate velocity.",
        style='LHM Body'
    )

    # Charts: Debt Trajectory and Interest Expense
    add_chart(doc, CHART_DIR / "15.png", "Figure 1.3: Federal Debt Trajectory — $36T debt, 124% of GDP")
    add_chart(doc, CHART_DIR / "16.png", "Figure 1.4: Federal Interest Expense — 11.3% of revenue and rising")

    doc.add_paragraph(
        "Meanwhile, the private components of GDP tell a very different story:",
        style='LHM Body'
    )
    doc.add_paragraph("• Private fixed investment is weak", style='LHM Body')
    doc.add_paragraph("• Goods consumption is flat to negative in real terms", style='LHM Body')
    doc.add_paragraph("• Net exports are a drag", style='LHM Body')
    doc.add_paragraph("• Capex outside of AI-linked megacaps is contracting", style='LHM Body')

    doc.add_paragraph(
        "This is not an expansion. It is levitation.",
        style='LHM Body'
    )

    add_page_break(doc)

    # Section III: GDI Tells the Truth
    doc.add_paragraph("III. GDI TELLS THE TRUTH", style='LHM H2')
    doc.add_paragraph("Private Income Has Stopped Growing", style='LHM Body')

    doc.add_paragraph(
        "GDI exposes what GDP conceals.",
        style='LHM Body'
    )

    doc.add_paragraph(
        "Real wage growth for the median worker has stalled. Nominal wages continue to rise, but not "
        "fast enough to offset cumulative inflation. Corporate profits are bifurcated: megacaps remain "
        "strong, while small and mid-sized firms face margin compression from labor, financing, and input costs.",
        style='LHM Body'
    )

    # Chart: Wealth Distribution
    add_chart(doc, CHART_DIR / "21.png", "Figure 1.5: Wealth Distribution — K-shaped reality in balance sheets")

    doc.add_paragraph(
        "Tax receipts at both the federal and state level have decelerated, confirming the slowdown "
        "in taxable income growth. Interest income has risen for asset-rich households but has not "
        "offset wage stagnation for the majority of consumers.",
        style='LHM Body'
    )

    doc.add_paragraph(
        "This is why sentiment surveys remain weak despite 'strong' headline data. Households experience "
        "the economy through income, not GDP.",
        style='LHM Body'
    )

    add_page_break(doc)

    # Section IV: Labor Market
    doc.add_paragraph("IV. LABOR: THE ECONOMY IS FREEZING, NOT FIRING", style='LHM H2')

    doc.add_paragraph(
        "The labor market is the most misunderstood component of the current cycle.",
        style='LHM Body'
    )

    doc.add_paragraph(
        "Headline metrics—non-farm payrolls and the unemployment rate—suggest resilience. "
        "Internals tell a very different story.",
        style='LHM Body'
    )

    doc.add_paragraph(
        "The defining feature of the 2025–2026 labor market is immobility.",
        style='LHM Body'
    )

    # Chart: Employment Diffusion
    add_chart(doc, CHART_DIR / "1.png", "Figure 1.6: Employment Diffusion Index — At neutral/contraction boundary")

    doc.add_paragraph(
        "The quits rate has collapsed below pre-pandemic levels. Temporary help employment has "
        "contracted sharply. Average weekly hours are at cycle lows. Continuing claims are rising "
        "steadily from the floor.",
        style='LHM Body'
    )

    # Chart: Labor Fragility Index
    add_chart(doc, CHART_DIR / "7.png", "Figure 1.7: Labor Fragility Index — Back to late-2019 warning levels")

    doc.add_paragraph(
        "These are not signals of a tight labor market. They are signals of a risk-averse labor market.",
        style='LHM Body'
    )

    doc.add_paragraph(
        "Workers are not quitting because they do not believe they can find better jobs. Employers "
        "are not hiring aggressively because they do not see sustained demand. Instead of layoffs, "
        "firms are cutting hours, freezing headcount, and deferring expansion.",
        style='LHM Body'
    )

    # Chart: LFI Components
    add_chart(doc, CHART_DIR / "19.png", "Figure 1.8: LFI Component Breakdown — Quits, hires/quits, long-duration")

    doc.add_paragraph(
        "This creates the false impression of stability.",
        style='LHM Body'
    )

    # Chart: Job Cuts vs Claims
    add_chart(doc, CHART_DIR / "32.png", "Figure 1.9: Job Cuts vs Initial Claims — Cuts at 2008 levels, claims suppressed")

    doc.add_paragraph(
        "Historically, layoffs rise after this phase, not during it. The freeze precedes the break.",
        style='LHM Body'
    )

    add_page_break(doc)

    # Section V: Consumer
    doc.add_paragraph("V. THE CONSUMER: FROM LIQUIDITY TO SOLVENCY", style='LHM H2')

    doc.add_paragraph(
        "For three years, consumer resilience was explained by excess savings. "
        "That explanation is now obsolete.",
        style='LHM Body'
    )

    # Chart: Excess Savings
    add_chart(doc, CHART_DIR / "33.png", "Figure 1.10: Excess Savings Depletion — Bottom 80% exhausted June 2023")

    doc.add_paragraph(
        "Excess savings are gone for the bottom 80% of households. Personal savings rates are below "
        "pre-pandemic norms. Credit card balances are at record highs. Delinquencies are accelerating, "
        "particularly in subprime and near-prime cohorts.",
        style='LHM Body'
    )

    doc.add_paragraph(
        "This is no longer a liquidity story. It is a solvency story.",
        style='LHM Body'
    )

    # Chart: Consumer Credit vs Savings
    add_chart(doc, CHART_DIR / "2.png", "Figure 1.11: Consumer Credit vs Savings — Rate at 4.0% vs 8.5% average")

    doc.add_paragraph(
        "Consumers are not spending because they are confident. They are spending because fixed "
        "expenses—rent, insurance, utilities—continue to rise, and discretionary adjustment has limits. "
        "Credit is filling the gap.",
        style='LHM Body'
    )

    # Chart: Two-Speed Consumer
    add_chart(doc, CHART_DIR / "22.png", "Figure 1.12: Two-Speed Consumer — Bifurcation in spending patterns")

    # Chart: Subprime Auto Delinquencies
    add_chart(doc, CHART_DIR / "31.png", "Figure 1.13: Subprime Auto Delinquencies — 6.67% exceeds 2008 peak")

    doc.add_paragraph(
        "This is why retail sales can remain positive even as financial stress builds. "
        "It is also why defaults rise suddenly rather than gradually.",
        style='LHM Body'
    )

    add_page_break(doc)

    # Section VI: K-Shaped Cost of Capital
    doc.add_paragraph("VI. THE K-SHAPED COST OF CAPITAL", style='LHM H2')

    doc.add_paragraph(
        "Monetary policy has not tightened the economy evenly. It has segmented it.",
        style='LHM Body'
    )

    doc.add_paragraph(
        "Large corporations refinanced at historically low rates during 2020–2021. They now earn more "
        "on cash than they pay on debt. Higher policy rates have improved their net interest income.",
        style='LHM Body'
    )

    doc.add_paragraph(
        "Small businesses did not have that option.",
        style='LHM Body'
    )

    doc.add_paragraph(
        "SMEs rely on floating-rate bank credit, private lending, and short-term financing. Their "
        "effective borrowing costs are now in the high single digits or higher. For these firms, "
        "higher rates are not a headwind—they are a choke point.",
        style='LHM Body'
    )

    # Chart: CRE Delinquencies
    add_chart(doc, CHART_DIR / "20.png", "Figure 1.14: Commercial Real Estate Delinquencies — Office at 11.76% exceeds 2008")

    doc.add_paragraph(
        "This bifurcation explains why equity indices dominated by megacaps appear strong while "
        "small caps, private investment, and hiring lag.",
        style='LHM Body'
    )

    doc.add_paragraph(
        "Markets are pricing the top decile of balance sheets and extrapolating that strength onto "
        "the entire economy. That extrapolation is incorrect.",
        style='LHM Body'
    )

    add_page_break(doc)

    # Section VII: Defining the Silent Stop
    doc.add_paragraph("VII. DEFINING THE SILENT STOP", style='LHM H2')

    doc.add_paragraph(
        "We define The Silent Stop as an economic condition where:",
        style='LHM Body'
    )

    doc.add_paragraph("• Headline growth remains positive due to fiscal flow", style='LHM Body')
    doc.add_paragraph("• Private income growth is flat or negative", style='LHM Body')
    doc.add_paragraph("• Labor mobility collapses before unemployment rises", style='LHM Body')
    doc.add_paragraph("• Consumer spending is maintained through credit, not income", style='LHM Body')
    doc.add_paragraph("• Capital formation outside large incumbents stalls", style='LHM Body')

    doc.add_paragraph(
        "This is not a recession in the traditional sense. It is a pre-recessionary stall that can "
        "persist longer than expected and end abruptly.",
        style='LHM Body'
    )

    doc.add_paragraph(
        "The danger is not that markets are wrong today. The danger is that they are pricing "
        "continuity in a regime defined by fragility.",
        style='LHM Body'
    )

    add_horizontal_line(doc)

    # Section VIII: Why This Matters
    doc.add_paragraph("VIII. WHY THIS MATTERS", style='LHM H2')

    doc.add_paragraph(
        "Stall regimes do not resolve gently.",
        style='LHM Body'
    )

    doc.add_paragraph(
        "They resolve when one of three things breaks:",
        style='LHM Body'
    )

    doc.add_paragraph("1. Funding transmission", style='LHM Body')
    doc.add_paragraph("2. Labor confidence", style='LHM Body')
    doc.add_paragraph("3. Credit availability", style='LHM Body')

    doc.add_paragraph(
        "Parts II–III will show why the system is now vulnerable on all three fronts simultaneously.",
        style='LHM Body'
    )

    add_page_break(doc)


def add_part2(doc):
    """Add Part II - Monetary Mechanics."""
    # Part header
    p = doc.add_paragraph("PART II", style='LHM H1')
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    p = doc.add_paragraph("MONETARY MECHANICS")
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.runs[0].font.name = 'Arial'
    p.runs[0].font.size = Pt(24)
    p.runs[0].font.bold = True
    p.runs[0].font.color.rgb = OCEAN_BLUE

    p = doc.add_paragraph("Reserve Scarcity & The 'Not-QE' Regime")
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.runs[0].font.name = 'Georgia'
    p.runs[0].font.size = Pt(12)
    p.runs[0].font.italic = True
    p.runs[0].font.color.rgb = NEUTRAL_GRAY

    add_horizontal_line(doc)

    doc.add_paragraph(
        '"Rates tell you intent. Reserves tell you constraint."',
        style='LHM Quote'
    )

    doc.add_paragraph(
        "Part I established the macro reality: the private economy has stalled beneath fiscal optics.",
        style='LHM Body'
    )

    doc.add_paragraph(
        "Part II explains how that stall was masked, why the masking mechanism is now exhausted, "
        "and why the Federal Reserve has already crossed the line from tightening to covert stabilization.",
        style='LHM Body'
    )

    doc.add_paragraph(
        "This is not a story about the fed funds rate. It is a story about settlement capacity, "
        "balance sheet constraints, and the quiet end of QT.",
        style='LHM Body'
    )

    add_horizontal_line(doc)

    # Section I: Liquidity Defined
    doc.add_paragraph("I. LIQUIDITY DEFINED CORRECTLY", style='LHM H2')

    doc.add_paragraph(
        "Liquidity is not 'risk-on.' Liquidity is not 'financial conditions.' "
        "Liquidity is not the level of asset prices.",
        style='LHM Body'
    )

    doc.add_paragraph(
        "Liquidity is the system's ability to clear cash and collateral at scale without discontinuity.",
        style='LHM Body'
    )

    doc.add_paragraph("Operationally, liquidity answers four questions:", style='LHM Body')
    doc.add_paragraph("1. Can banks meet payment and regulatory needs without hoarding reserves?", style='LHM Body')
    doc.add_paragraph("2. Can dealers intermediate Treasury supply without balance sheet stress?", style='LHM Body')
    doc.add_paragraph("3. Can money funds place cash without distorting overnight rates?", style='LHM Body')
    doc.add_paragraph("4. Can Treasury issue debt without draining reserves non-linearly?", style='LHM Body')

    doc.add_paragraph(
        "When the answer to any of these becomes 'no,' markets can break without any macro catalyst.",
        style='LHM Body'
    )

    # New Priority 1 Chart: TGA Balance
    p1_tga = PRIORITY1_DIR / "chart_tga_balance.png"
    if p1_tga.exists():
        add_chart(doc, p1_tga, "Figure 2.0: TGA Balance History — Treasury cash volatility as liquidity lever")

    add_page_break(doc)

    # Section: Reserve Scarcity
    doc.add_paragraph("II. BANK RESERVES: THE ADEQUACY QUESTION", style='LHM H2')

    doc.add_paragraph(
        "Reserve adequacy is not about the absolute level. It is about distribution, behavior, "
        "and willingness to deploy.",
        style='LHM Body'
    )

    # Chart: Bank Reserves vs GDP
    add_chart(doc, CHART_DIR / "8.png", "Figure 2.1: Bank Reserves vs GDP — 12.5%, approaching Sept 2019 crisis level")

    doc.add_paragraph(
        "Reserves become scarce when they are concentrated in a small set of banks, regulatory "
        "constraints raise the cost of intermediation, dealers refuse to expand balance sheets, "
        "and banks hoard reserves defensively.",
        style='LHM Body'
    )

    doc.add_paragraph(
        "In 2019, repo markets broke with reserves well above $1T. Today, the system is more "
        "constrained than 2019, not less: Higher G-SIB surcharges, lower dealer balance sheet "
        "elasticity, far greater Treasury supply.",
        style='LHM Body'
    )

    # New Priority 1 Chart: Fed Balance Sheet
    p1_fed = PRIORITY1_DIR / "chart_fed_balance_sheet.png"
    if p1_fed.exists():
        add_chart(doc, p1_fed, "Figure 2.2: Fed Balance Sheet Composition — RRP offset now exhausted")

    add_page_break(doc)

    # Section: Treasury Supply Challenge
    doc.add_paragraph("III. THE TREASURY SUPPLY CHALLENGE", style='LHM H2')

    doc.add_paragraph(
        "With RRP gone, the Treasury General Account becomes a direct liquidity weapon. "
        "TGA up means reserves down.",
        style='LHM Body'
    )

    # Chart: Treasury Maturity Wall
    add_chart(doc, CHART_DIR / "25.png", "Figure 2.3: Treasury Maturity Wall — $12.7T in 36 months (3x average)")

    doc.add_paragraph(
        "Treasury cash management is now de facto monetary policy. This is why debt ceiling dynamics, "
        "refunding schedules, and issuance composition now matter far more than rate rhetoric.",
        style='LHM Body'
    )

    # Chart: Treasury Issuance by Tenor
    add_chart(doc, CHART_DIR / "31.png", "Figure 2.4: Treasury Issuance by Tenor")

    add_page_break(doc)

    # Section: Funding Market Stress
    doc.add_paragraph("IV. FUNDING MARKET STRESS SIGNALS", style='LHM H2')

    doc.add_paragraph(
        "In reserve-constrained regimes, small spreads matter. EFFR hugs IORB. SOFR trades rich. "
        "These are not noise. They are the earliest indicators of binding constraints.",
        style='LHM Body'
    )

    # Chart: SOFR-EFFR Spread
    add_chart(doc, CHART_DIR / "10.png", "Figure 2.5: SOFR-EFFR Spread — 6.1 bps in warning zone")

    doc.add_paragraph(
        "Plumbing breaks do not start with headlines. They start with basis points.",
        style='LHM Body'
    )

    # Chart: Repo Dispersion
    add_chart(doc, CHART_DIR / "4.png", "Figure 2.6: Repo Dispersion — Widening indicates dealer stress")

    add_page_break(doc)

    # Section: SRF
    doc.add_paragraph("V. THE STANDING REPO FACILITY", style='LHM H2')

    doc.add_paragraph(
        "The SRF is irrelevant in abundant liquidity regimes. It becomes relevant only when "
        "private intermediation fails.",
        style='LHM Body'
    )

    # Chart: SRF Usage
    add_chart(doc, CHART_DIR / "26.png", "Figure 2.7: Standing Repo Facility Usage — Record $48.2B October 31, 2025")

    doc.add_paragraph(
        "Rising or persistent SRF usage tells you dealer balance sheets are constrained, "
        "collateral velocity is impaired, and the system is testing the ceiling.",
        style='LHM Body'
    )

    doc.add_paragraph(
        "When both floor tools (RRP) and ceiling tools (SRF) are implicated in the same cycle, "
        "the system is operating at the edge.",
        style='LHM Body'
    )

    add_page_break(doc)

    # Section: Dealer Balance Sheets
    doc.add_paragraph("VI. DEALER BALANCE SHEET CONSTRAINTS", style='LHM H2')

    # Chart: Primary Dealer Balance Sheet
    add_chart(doc, CHART_DIR / "5.png", "Figure 2.8: Primary Dealer Balance Sheet — $94B inventory, 78% of SLR")

    doc.add_paragraph(
        "Dealers are already running at 78% of their Supplementary Leverage Ratio capacity. "
        "Only $26B of headroom remains before regulatory constraints become binding.",
        style='LHM Body'
    )

    # Chart: Treasury Basis
    add_chart(doc, CHART_DIR / "19.png", "Figure 2.9: Treasury Basis Dynamics — -27.9 bps below unwind threshold")

    add_page_break(doc)

    # Section: RMPs
    doc.add_paragraph("VII. RESERVE MANAGEMENT PURCHASES (RMPs)", style='LHM H2')
    doc.add_paragraph("When Implementation Becomes Policy", style='LHM Body')

    doc.add_paragraph(
        '"RMPs exist because QT hit the wall."',
        style='LHM Quote'
    )

    doc.add_paragraph(
        "Reserve Management Purchases are outright Treasury purchases conducted explicitly to add reserves. "
        "They are framed as technical, temporary, implementation-focused, and not accommodative.",
        style='LHM Body'
    )

    doc.add_paragraph(
        "That framing is rhetorical.",
        style='LHM Body'
    )

    doc.add_paragraph(
        "Mechanically: RMPs increase reserve balances. Period. They are QT offsets.",
        style='LHM Body'
    )

    doc.add_paragraph(
        "RMPs appear only when the Fed believes reserve scarcity threatens rate control. Their emergence "
        "tells you three things simultaneously:",
        style='LHM Body'
    )
    doc.add_paragraph("1. QT can no longer proceed unbuffered", style='LHM Body')
    doc.add_paragraph("2. The reserve floor is near", style='LHM Body')
    doc.add_paragraph("3. Funding stress risk is rising faster than policy can respond", style='LHM Body')

    doc.add_paragraph(
        "If QT were harmless, RMPs would not exist.",
        style='LHM Body'
    )

    add_horizontal_line(doc)

    doc.add_paragraph("VIII. WHAT PART II PROVES", style='LHM H2')

    doc.add_paragraph(
        "Part II proves one thing conclusively:",
        style='LHM Body'
    )

    doc.add_paragraph(
        '"The Fed has already acknowledged reserve scarcity—through its actions, not its words."',
        style='LHM Quote'
    )

    doc.add_paragraph(
        "RMPs are the tell. They are the quiet admission that QT hit the wall. "
        "Everything downstream—volatility, correlation failure, credit repricing—follows from that fact.",
        style='LHM Body'
    )

    add_page_break(doc)


def add_part3(doc):
    """Add Part III - Market Technicals."""
    # Part header
    p = doc.add_paragraph("PART III", style='LHM H1')
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    p = doc.add_paragraph("MARKET TECHNICALS")
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.runs[0].font.name = 'Arial'
    p.runs[0].font.size = Pt(24)
    p.runs[0].font.bold = True
    p.runs[0].font.color.rgb = OCEAN_BLUE

    p = doc.add_paragraph("Liquidity-Driven Price Action & Positioning")
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.runs[0].font.name = 'Georgia'
    p.runs[0].font.size = Pt(12)
    p.runs[0].font.italic = True
    p.runs[0].font.color.rgb = NEUTRAL_GRAY

    add_horizontal_line(doc)

    doc.add_paragraph(
        '"When reserves are the constraint, markets stop trending and start snapping."',
        style='LHM Quote'
    )

    # Section I: Plumbing to Price
    doc.add_paragraph("I. FROM PLUMBING TO PRICE ACTION", style='LHM H2')

    doc.add_paragraph(
        "Once the Fed is forced into RMPs, the regime shifts in three important ways:",
        style='LHM Body'
    )

    doc.add_paragraph(
        "1. Liquidity becomes endogenous — Markets are no longer responding primarily to growth or "
        "inflation data. They are responding to whether reserves are being added fast enough.",
        style='LHM Body'
    )

    doc.add_paragraph(
        "2. Price action decouples from fundamentals — Asset prices can remain elevated even as "
        "fundamentals deteriorate because marginal stabilization is being supplied mechanically.",
        style='LHM Body'
    )

    doc.add_paragraph(
        "3. Risk becomes discontinuous — When stabilization depends on timing and calibration, "
        "misalignment produces gaps rather than trends.",
        style='LHM Body'
    )

    add_page_break(doc)

    # Section: Treasury Auction Stress
    doc.add_paragraph("II. TREASURY AUCTION STRESS", style='LHM H2')

    # Chart: Auction Tails Time Series
    add_chart(doc, CHART_DIR / "9.png", "Figure 3.1: Auction Tails Time Series — Frequency increasing")

    # Chart: Auction Tails Scatter
    add_chart(doc, CHART_DIR / "10.png", "Figure 3.2: Auction Tails Scatter — Long-duration showing most stress")

    # Chart: Treasury Auction Deviation
    add_chart(doc, CHART_DIR / "34.png", "Figure 3.3: Treasury Auction Deviation Analysis")

    add_page_break(doc)

    # Section: Yield Curve Dynamics
    doc.add_paragraph("III. YIELD CURVE DYNAMICS", style='LHM H2')

    # Chart: Yield Curve Shape
    add_chart(doc, CHART_DIR / "16.png", "Figure 3.4: Treasury Yield Curve Shape Analysis")

    # Chart: Yield Curve Heatmap
    add_chart(doc, CHART_DIR / "8.png", "Figure 3.5: Yield Curve Heatmap — Historical configurations")

    # Chart: Repricing Path
    add_chart(doc, CHART_DIR / "22.png", "Figure 3.6: Yield Curve Repricing Path — 10Y to 5.10% (+95 bps)")

    # Chart: 10Y Scenario Analysis
    add_chart(doc, CHART_DIR / "23.png", "Figure 3.7: Ten-Year Yield Scenario Analysis")

    # New Priority 1 Chart: Real Rates
    p1_real = PRIORITY1_DIR / "chart_real_rates.png"
    if p1_real.exists():
        add_chart(doc, p1_real, "Figure 3.8: Real Rates & Breakevens — Inflation expectations regime")

    add_page_break(doc)

    # Section: Credit Market Analysis
    doc.add_paragraph("IV. CREDIT MARKET ANALYSIS", style='LHM H2')

    doc.add_paragraph(
        "Credit spreads remain tight because carry remains attractive and RMPs suppress systemic stress. "
        "But credit is downstream of reserves, not upstream.",
        style='LHM Body'
    )

    # Chart: Credit Spread Gauges
    add_chart(doc, CHART_DIR / "29.png", "Figure 3.9: Credit Spread Percentile Gauges — HY at 28th percentile")

    doc.add_paragraph(
        "When reserve scarcity tightens: refinancing becomes harder, dealer appetite for credit risk "
        "falls, ETF outflows accelerate, spreads gap rather than widen gradually.",
        style='LHM Body'
    )

    # Chart: Credit Spread Waterfall
    add_chart(doc, CHART_DIR / "27.png", "Figure 3.10: Credit Spread Waterfall — Decomposition")

    # Chart: Credit Impulse
    add_chart(doc, CHART_DIR / "26.png", "Figure 3.11: Credit Impulse — Leading indicator for real economy")

    doc.add_paragraph(
        "Credit should be treated as late-cycle convex risk, not yield pickup.",
        style='LHM Body'
    )

    add_page_break(doc)

    # Section: Cross-Asset Signals
    doc.add_paragraph("V. CROSS-ASSET SIGNALS", style='LHM H2')

    # Chart: Cross-Asset Correlations
    add_chart(doc, CHART_DIR / "33.png", "Figure 3.12: Cross-Asset Correlations — Stock-bond at -0.48")

    doc.add_paragraph(
        "One dangerous assumption investors carry into 2026: that duration will reliably hedge equity risk.",
        style='LHM Body'
    )

    doc.add_paragraph(
        "When term premium is rising, Treasury supply is heavy, and liquidity support is targeted—equities "
        "and bonds can sell off together. Correlation breakdown is a plumbing failure, not a macro surprise.",
        style='LHM Body'
    )

    # Chart: Foreign Official Holdings
    add_chart(doc, CHART_DIR / "30.png", "Figure 3.13: Foreign Official Treasury Holdings — China declining structurally")

    # New Priority 1 Chart: Dollar Index
    p1_dollar = PRIORITY1_DIR / "chart_dollar_index.png"
    if p1_dollar.exists():
        add_chart(doc, p1_dollar, "Figure 3.14: Trade-Weighted Dollar Index — USD funding demand")

    add_page_break(doc)

    # Section: Volatility Regime
    doc.add_paragraph("VI. VOLATILITY REGIME", style='LHM H2')

    doc.add_paragraph(
        "RMPs suppress volatility between stress windows, but they do not eliminate it. "
        "This produces a dangerous configuration: low spot volatility, shallow contango, cheap short-dated convexity.",
        style='LHM Body'
    )

    # New Priority 1 Chart: MOVE vs VIX
    p1_vol = PRIORITY1_DIR / "chart_move_vs_vix.png"
    if p1_vol.exists():
        add_chart(doc, p1_vol, "Figure 3.15: MOVE vs VIX — Rates volatility often leads equity volatility")

    doc.add_paragraph(
        "Historically, this is when volatility reprices violently—from a low base, in compressed time. "
        "In reserve-scarce regimes, volatility is not cyclical; it is episodic.",
        style='LHM Body'
    )

    add_horizontal_line(doc)

    # Synthesis
    doc.add_paragraph("VII. SYNTHESIS: HOW RMP REGIMES BREAK", style='LHM H2')

    doc.add_paragraph("The sequence:", style='LHM Body')
    doc.add_paragraph("1. RRP is gone", style='LHM Body')
    doc.add_paragraph("2. Reserves are scarce", style='LHM Body')
    doc.add_paragraph("3. QT is capped by RMPs", style='LHM Body')
    doc.add_paragraph("4. Markets extrapolate stability", style='LHM Body')
    doc.add_paragraph("5. Positioning accumulates", style='LHM Body')
    doc.add_paragraph("6. A funding or issuance shock hits", style='LHM Body')
    doc.add_paragraph("7. RMPs respond with lag", style='LHM Body')
    doc.add_paragraph("8. Liquidity evaporates", style='LHM Body')
    doc.add_paragraph("9. Markets gap", style='LHM Body')

    doc.add_paragraph(
        "The error investors make: assuming that because the Fed can stabilize the system, "
        "markets will adjust smoothly.",
        style='LHM Body'
    )

    doc.add_paragraph(
        "Stabilization in a reserve-scarce regime is discrete, reactive, and uneven.",
        style='LHM Body'
    )

    add_page_break(doc)


def add_conclusion(doc):
    """Add the Conclusion section."""
    # Conclusion header
    p = doc.add_paragraph("CONCLUSION", style='LHM H1')
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    p = doc.add_paragraph("THE SILENT STOP, REVISITED")
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.runs[0].font.name = 'Arial'
    p.runs[0].font.size = Pt(24)
    p.runs[0].font.bold = True
    p.runs[0].font.color.rgb = OCEAN_BLUE

    p = doc.add_paragraph("Why This Cycle Ends with Discontinuity, Not Resolution")
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.runs[0].font.name = 'Georgia'
    p.runs[0].font.size = Pt(12)
    p.runs[0].font.italic = True
    p.runs[0].font.color.rgb = NEUTRAL_GRAY

    add_horizontal_line(doc)

    # Chart: 16-Week Stress Window
    add_chart(doc, CHART_DIR / "11.png", "Figure 4.1: The 16-Week Stress Window — Critical event calendar")

    doc.add_paragraph(
        "The defining error of the current macro moment is not optimism. It is misdiagnosis.",
        style='LHM Body'
    )

    doc.add_paragraph(
        "The economy has not normalized. It has stalled.",
        style='LHM Body'
    )

    doc.add_paragraph(
        "GDI is flat. Labor mobility has collapsed. Hours are at cycle lows. The consumer has shifted "
        "to solvency stress. SMEs face prohibitive capital costs.",
        style='LHM Body'
    )

    doc.add_paragraph(
        "These are not late-cycle warnings. They are late-cycle facts.",
        style='LHM Body'
    )

    add_horizontal_line(doc)

    # Core Synthesis
    doc.add_paragraph("THE CORE SYNTHESIS", style='LHM H2')

    doc.add_paragraph("1. THIS IS NOT A SLOWDOWN — IT IS A LOSS OF VELOCITY", style='LHM H2')
    doc.add_paragraph(
        "In a normal slowdown, the economy moves forward more slowly. In this regime, motion itself "
        "has broken down. The labor market is freezing. The consumer is substituting credit for income. "
        "Economic velocity has fallen toward zero.",
        style='LHM Body'
    )

    doc.add_paragraph("2. FISCAL SUPPORT IS MASKING, NOT FIXING", style='LHM H2')
    doc.add_paragraph(
        "Government spending has prevented contraction. It has not restored private-sector momentum. "
        "When fiscal impulse fades, the private economy will not reaccelerate. It will be exposed.",
        style='LHM Body'
    )

    doc.add_paragraph("3. LIQUIDITY HAS BECOME A CONSTRAINT", style='LHM H2')
    doc.add_paragraph(
        "The liquidity buffer has been exhausted. Reserve scarcity is now the binding constraint. "
        "RMPs are an admission that QT reached its limit. Liquidity support is now conditional, "
        "reactive, and uneven.",
        style='LHM Body'
    )

    doc.add_paragraph(
        "In such regimes: Volatility is suppressed until it isn't. Correlations break when needed most. "
        "Credit reprices abruptly. Markets gap instead of trend.",
        style='LHM Body'
    )

    doc.add_paragraph("4. MARKETS ARE PRICING CONTINUITY IN A FRAGILE SYSTEM", style='LHM H2')
    doc.add_paragraph(
        "Asset prices assume the current equilibrium can persist. Market internals tell a different story. "
        "History suggests such regimes end when one transmission channel fails: Funding, Credit, or Confidence.",
        style='LHM Body'
    )

    add_page_break(doc)

    # What This Means
    doc.add_paragraph("WHAT THIS MEANS FOR INVESTORS", style='LHM H2')

    doc.add_paragraph(
        "The dominant risk is not being early on a slowdown. It is being structurally exposed to "
        "nonlinear repricing.",
        style='LHM Body'
    )

    doc.add_paragraph("In this environment:", style='LHM Body')
    doc.add_paragraph("• Carry trades work until they don't", style='LHM Body')
    doc.add_paragraph("• Diversification fails intermittently", style='LHM Body')
    doc.add_paragraph("• Liquidity disappears when most needed", style='LHM Body')
    doc.add_paragraph("• The cost of being wrong is asymmetric", style='LHM Body')

    doc.add_paragraph(
        "Portfolios should be built to survive gaps, not optimize for smooth paths.",
        style='LHM Body'
    )

    doc.add_paragraph(
        "The goal is not to avoid risk. It is to avoid being forced.",
        style='LHM Body'
    )

    add_horizontal_line(doc)

    # Final Thought
    doc.add_paragraph("FINAL THOUGHT", style='LHM H2')

    doc.add_paragraph(
        "This is not a soft landing. It is not a hard landing. It is not a reacceleration.",
        style='LHM Body'
    )

    doc.add_paragraph(
        "It is levitation over a vacuum.",
        style='LHM Body'
    )

    doc.add_paragraph(
        "The economy has stopped moving forward, but asset prices have not yet reconciled with that reality.",
        style='LHM Body'
    )

    doc.add_paragraph(
        "That gap will close. When it does, the adjustment will not be linear.",
        style='LHM Body'
    )

    doc.add_paragraph(
        "The Silent Stop does not announce itself. It reveals itself—all at once.",
        style='LHM Body'
    )

    doc.add_paragraph(
        "Position accordingly.",
        style='LHM Body'
    )

    add_horizontal_line(doc)

    # Sign-off
    for _ in range(2):
        doc.add_paragraph()

    p = doc.add_paragraph(
        "That's our view from the Watch. Until next time, we'll be sure to keep the light on...."
    )
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.runs[0].font.name = 'Georgia'
    p.runs[0].font.size = Pt(11)
    p.runs[0].font.italic = True
    p.runs[0].font.color.rgb = NEUTRAL_GRAY

    for _ in range(2):
        doc.add_paragraph()

    p = doc.add_paragraph("— Bob Sheehan, CFA, CMT")
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.runs[0].font.name = 'Arial'
    p.runs[0].font.size = Pt(12)
    p.runs[0].font.bold = True

    p = doc.add_paragraph("Founder & Chief Investment Officer")
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.runs[0].font.name = 'Arial'
    p.runs[0].font.size = Pt(11)

    p = doc.add_paragraph("Lighthouse Macro")
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.runs[0].font.name = 'Arial'
    p.runs[0].font.size = Pt(11)
    p.runs[0].font.color.rgb = OCEAN_BLUE

    for _ in range(2):
        doc.add_paragraph()

    p = doc.add_paragraph("MACRO, ILLUMINATED.")
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.runs[0].font.name = 'Arial'
    p.runs[0].font.size = Pt(14)
    p.runs[0].font.bold = True
    p.runs[0].font.color.rgb = OCEAN_BLUE


def generate_report():
    """Generate the complete Horizon Report."""
    print("=" * 60)
    print("GENERATING THE HORIZON REPORT - JANUARY 2026")
    print("=" * 60)

    # Create document
    print("\nCreating document structure...")
    doc = create_document()

    # Add sections
    print("Adding cover page...")
    add_cover_page(doc)

    print("Adding table of contents...")
    add_toc(doc)

    print("Adding executive summary...")
    add_executive_summary(doc)

    print("Adding Part I - The Silent Stop...")
    add_part1(doc)

    print("Adding Part II - Monetary Mechanics...")
    add_part2(doc)

    print("Adding Part III - Market Technicals...")
    add_part3(doc)

    print("Adding Conclusion...")
    add_conclusion(doc)

    # Save
    timestamp = datetime.now().strftime("%Y%m%d_%H%M")
    output_path = OUTPUT_DIR / f"The_Horizon_Report_January_2026_FINAL_{timestamp}.docx"

    print(f"\nSaving to: {output_path}")
    doc.save(str(output_path))

    # Also save as "latest"
    latest_path = OUTPUT_DIR / "The_Horizon_Report_January_2026_FINAL.docx"
    doc.save(str(latest_path))

    print("\n" + "=" * 60)
    print("REPORT GENERATION COMPLETE")
    print("=" * 60)
    print(f"\nOutput: {output_path}")
    print(f"Latest: {latest_path}")

    return output_path


if __name__ == "__main__":
    generate_report()
