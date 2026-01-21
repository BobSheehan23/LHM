"""
The Horizon Report - January 2026
Final Institutional-Quality Report Generator V2

Uses the updated 36 Horizon charts + 7 Priority 1 charts
Lighthouse Macro - MACRO, ILLUMINATED.
"""

from docx import Document
from docx.shared import Inches, Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from pathlib import Path
from datetime import datetime

# Paths
CHART_DIR = Path("/Users/bob/LHM/deliverables/horizon_report_final/charts")
PRIORITY1_DIR = Path("/Users/bob/LHM/data/charts/priority1")
OUTPUT_DIR = Path("/Users/bob/LHM/deliverables/horizon_report_final")

# Chart mapping from document specification
CHART_MAP = {
    # Part I - The Silent Stop
    "1.1": "27_Real_vs_Nominal_GDP.png",
    "1.2": "12_Fiscal_Dominance_Cascade.png",
    "1.3": "15_Federal_Debt_Trajectory.png",
    "1.4": "16_Federal_Interest_Expense.png",
    "1.5": "21_Wealth_Distribution_and_Balance_Sheet_Divergence.png",
    "1.6": "01_Employment_Diffusion_Index.png",
    "1.7": "07_Labor_Fragility_Index.png",
    "1.8": "19_Labor_Fragility_Index.png",
    "1.9": "32_Job_Cuts_vs_Initial_Claims.png",
    "1.10": "33_Excess_Savings_Alternative_View.png",
    "1.11": "02_Consumer_Credit_vs_Savings.png",
    "1.12": "22_Two_Speed_Consumer.png",
    "1.13": "31_Subprime_Auto_Delinquencies.png",
    "1.14": "20_Commercial_Real_Estate_Delinquencies.png",

    # Part II - Monetary Mechanics
    "2.1": "08_Bank_Reserves_vs_GDP.png",
    "2.2": "25_Treasury_Maturity_Wall.png",
    "2.3": "31_Treasury_Issuance_by_Tenor.png",
    "2.4": "10_SOFR_minus_EFFR_Spread.png",
    "2.5": "04_Repo_Dispersion.png",
    "2.6": "26_Standing_Repo_Facility_Usage.png",
    "2.7": "02_SRF_Usage.png",
    "2.8": "05_Primary_Dealer_Balance_Sheet.png",
    "2.9": "19_Treasury_Basis_Dynamics.png",

    # Part III - Market Technicals
    "3.1": "09_Auction_Tails_Time_Series.png",
    "3.2": "10_Auction_Tails_Scatter.png",
    "3.3": "34_Treasury_Auction_Tails_Deviation_Analysis.png",
    "3.4": "16_Treasury_Yield_Curve_Shape_Analysis.png",
    "3.5": "08_Yield_Curve_Heatmap.png",
    "3.6": "22_Yield_Curve_Repricing_Path.png",
    "3.7": "23_Ten_Year_Yield_Scenario_Analysis.png",
    "3.8": "29_Credit_Spread_Distribution_Gauges.png",
    "3.9": "27_Credit_Spread_Waterfall.png",
    "3.10": "26_Credit_Impulse.png",
    "3.11": "33_Cross_Asset_Correlations.png",
    "3.12": "30_Foreign_Official_Treasury_Holdings.png",

    # Conclusion
    "4.1": "11_Critical_Event_Calendar_16_Week_Stress_Window.png",
}

# Lighthouse Colors
OCEAN_BLUE = RGBColor(0x00, 0x89, 0xD1)
DUSK_ORANGE = RGBColor(0xFF, 0x67, 0x23)
NEUTRAL_GRAY = RGBColor(0x80, 0x80, 0x80)
BLACK = RGBColor(0x00, 0x00, 0x00)
WHITE = RGBColor(0xFF, 0xFF, 0xFF)


def set_cell_shading(cell, color_hex):
    """Set cell background color."""
    shading_elm = OxmlElement('w:shd')
    shading_elm.set(qn('w:fill'), color_hex)
    cell._tc.get_or_add_tcPr().append(shading_elm)


def add_horizontal_line(doc):
    """Add a styled horizontal line."""
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after = Pt(12)
    run = p.add_run("─" * 85)
    run.font.color.rgb = NEUTRAL_GRAY
    run.font.size = Pt(8)


def add_page_break(doc):
    """Add page break."""
    doc.add_page_break()


def get_chart_path(figure_num):
    """Get the path to a chart by figure number."""
    if figure_num in CHART_MAP:
        return CHART_DIR / CHART_MAP[figure_num]
    return None


def add_chart(doc, figure_num, caption=None, width=6.3):
    """Add a chart image with caption."""
    chart_path = get_chart_path(figure_num)

    if chart_path and chart_path.exists():
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run()
        run.add_picture(str(chart_path), width=Inches(width))

        if caption:
            cap = doc.add_paragraph()
            cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
            cap_run = cap.add_run(f"Figure {figure_num}: {caption}")
            cap_run.font.size = Pt(9)
            cap_run.font.italic = True
            cap_run.font.color.rgb = NEUTRAL_GRAY
            cap.paragraph_format.space_after = Pt(12)
        return True
    else:
        # Placeholder for missing chart
        p = doc.add_paragraph(f"[Figure {figure_num}: {caption or 'Chart'}]")
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.runs[0].font.color.rgb = DUSK_ORANGE
        p.runs[0].font.italic = True
        return False


def add_priority1_chart(doc, chart_name, caption, width=6.3):
    """Add a Priority 1 chart."""
    chart_path = PRIORITY1_DIR / chart_name

    if chart_path.exists():
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run()
        run.add_picture(str(chart_path), width=Inches(width))

        cap = doc.add_paragraph()
        cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
        cap_run = cap.add_run(caption)
        cap_run.font.size = Pt(9)
        cap_run.font.italic = True
        cap_run.font.color.rgb = NEUTRAL_GRAY
        return True
    return False


def create_document():
    """Create the document with Lighthouse styling."""
    doc = Document()

    # Set margins
    for section in doc.sections:
        section.top_margin = Inches(0.7)
        section.bottom_margin = Inches(0.7)
        section.left_margin = Inches(0.9)
        section.right_margin = Inches(0.9)

    return doc


def add_styled_paragraph(doc, text, font_name='Georgia', font_size=11, bold=False,
                         italic=False, color=None, alignment=None, space_after=8):
    """Add a styled paragraph."""
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.name = font_name
    run.font.size = Pt(font_size)
    run.font.bold = bold
    run.font.italic = italic
    if color:
        run.font.color.rgb = color
    if alignment:
        p.alignment = alignment
    p.paragraph_format.space_after = Pt(space_after)
    p.paragraph_format.line_spacing = 1.15
    return p


def add_quote(doc, text):
    """Add a styled quote."""
    p = doc.add_paragraph()
    run = p.add_run(f'"{text}"')
    run.font.name = 'Georgia'
    run.font.size = Pt(12)
    run.font.italic = True
    run.font.color.rgb = OCEAN_BLUE
    p.paragraph_format.left_indent = Inches(0.4)
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after = Pt(12)
    return p


def add_section_header(doc, text, level=1):
    """Add a section header."""
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.name = 'Arial'

    if level == 1:
        run.font.size = Pt(18)
        run.font.bold = True
        run.font.color.rgb = OCEAN_BLUE
        p.paragraph_format.space_before = Pt(24)
        p.paragraph_format.space_after = Pt(12)
    else:
        run.font.size = Pt(14)
        run.font.bold = True
        run.font.color.rgb = BLACK
        p.paragraph_format.space_before = Pt(18)
        p.paragraph_format.space_after = Pt(6)

    return p


def add_bullet(doc, text):
    """Add a bullet point."""
    p = doc.add_paragraph(style='List Bullet')
    run = p.add_run(text)
    run.font.name = 'Georgia'
    run.font.size = Pt(11)
    p.paragraph_format.space_after = Pt(4)
    return p


def add_cover_page(doc):
    """Add the cover page."""
    for _ in range(4):
        doc.add_paragraph()

    # Brand
    add_styled_paragraph(doc, "LIGHTHOUSE MACRO", 'Arial', 14, bold=True,
                        color=OCEAN_BLUE, alignment=WD_ALIGN_PARAGRAPH.CENTER)
    add_styled_paragraph(doc, "MACRO, ILLUMINATED.", 'Arial', 10,
                        color=NEUTRAL_GRAY, alignment=WD_ALIGN_PARAGRAPH.CENTER)

    add_horizontal_line(doc)

    for _ in range(2):
        doc.add_paragraph()

    # Title
    add_styled_paragraph(doc, "THE HORIZON REPORT", 'Arial', 36, bold=True,
                        color=OCEAN_BLUE, alignment=WD_ALIGN_PARAGRAPH.CENTER)
    add_styled_paragraph(doc, "JANUARY 2026", 'Arial', 18,
                        color=NEUTRAL_GRAY, alignment=WD_ALIGN_PARAGRAPH.CENTER)

    add_horizontal_line(doc)

    # Subtitle
    add_styled_paragraph(doc, "THE SILENT STOP", 'Arial', 24, bold=True,
                        color=BLACK, alignment=WD_ALIGN_PARAGRAPH.CENTER)
    add_styled_paragraph(doc, "Why GDP Is Lying, Liquidity Is Constrained, and Markets Are Fragile",
                        'Georgia', 14, italic=True, color=NEUTRAL_GRAY,
                        alignment=WD_ALIGN_PARAGRAPH.CENTER)

    for _ in range(6):
        doc.add_paragraph()

    # Author
    add_styled_paragraph(doc, "Robert Sheehan, CFA, CMT", 'Arial', 12, bold=True,
                        alignment=WD_ALIGN_PARAGRAPH.CENTER)
    add_styled_paragraph(doc, "Founder & Chief Investment Officer", 'Arial', 11,
                        color=NEUTRAL_GRAY, alignment=WD_ALIGN_PARAGRAPH.CENTER)

    add_page_break(doc)


def add_executive_summary(doc):
    """Add executive summary."""
    add_section_header(doc, "EXECUTIVE SUMMARY")
    add_horizontal_line(doc)

    add_section_header(doc, "THE CORE THESIS: THE SILENT STOP", level=2)

    add_styled_paragraph(doc,
        "The U.S. economy in early 2026 is not overheating. It is stalling silently. "
        "Headline GDP remains positive (+1.8% YoY), but this figure is propped up entirely "
        "by fiscal injections. The private sector engine—measured by Gross Domestic Income—has "
        "flatlined at 0.0% YoY.")

    add_styled_paragraph(doc,
        "This is the largest GDP–GDI divergence on record. Historically, such dislocations "
        "resolve via downward GDP revisions—not upward GDI rebounds.")

    add_quote(doc, "GDP is the paint job. GDI is the engine. And the engine has seized.")

    # Three pillars
    add_section_header(doc, "THE THREE PILLARS OF FRAGILITY", level=2)

    # Pillar 1
    add_styled_paragraph(doc, "1. MACRO DYNAMICS: The Big Freeze", 'Arial', 12, bold=True, color=OCEAN_BLUE)
    add_styled_paragraph(doc, "The labor market appears stable on the surface, but internal dynamism has collapsed:")

    table = doc.add_table(rows=5, cols=2)
    table.style = 'Table Grid'
    headers = ["Indicator", "Current Reading"]
    data = [
        ("Quits Rate", "1.8% (below 2019 levels)"),
        ("Temp Help Employment (YoY)", "−15% (recession-consistent)"),
        ("Avg Weekly Hours", "34.2 (cycle low)"),
        ("Continuing Claims", "~1.97M (rising trend)"),
    ]

    for j, header in enumerate(headers):
        cell = table.rows[0].cells[j]
        cell.text = header
        cell.paragraphs[0].runs[0].font.bold = True
        set_cell_shading(cell, "0089D1")
        cell.paragraphs[0].runs[0].font.color.rgb = WHITE

    for i, (col1, col2) in enumerate(data, 1):
        table.rows[i].cells[0].text = col1
        table.rows[i].cells[1].text = col2

    doc.add_paragraph()
    add_styled_paragraph(doc, "The economy isn't laying off. It's locking up.", italic=True)

    # Pillar 2
    add_styled_paragraph(doc, "2. MONETARY MECHANICS: The Plumbing Break", 'Arial', 12, bold=True, color=OCEAN_BLUE)
    add_styled_paragraph(doc,
        "The system's shock absorber has been drained. The Overnight Reverse Repo Facility (RRP), "
        "which absorbed over $2 trillion in excess liquidity in 2022, is now effectively empty.")

    table2 = doc.add_table(rows=5, cols=2)
    table2.style = 'Table Grid'
    data2 = [
        ("Metric", "Current Reading"),
        ("RRP Balance", "$4.58B (from $2T+ in 2022)"),
        ("Reserve Balances", "~$2.96T (≈$180B above floor)"),
        ("SOFR vs IORB", "SOFR above IORB (stress signal)"),
        ("SRF Usage (Year-End)", "Record levels"),
    ]

    for i, (col1, col2) in enumerate(data2):
        table2.rows[i].cells[0].text = col1
        table2.rows[i].cells[1].text = col2
        if i == 0:
            for cell in table2.rows[i].cells:
                cell.paragraphs[0].runs[0].font.bold = True
                set_cell_shading(cell, "0089D1")
                cell.paragraphs[0].runs[0].font.color.rgb = WHITE

    doc.add_paragraph()
    add_styled_paragraph(doc,
        "The Fed has quietly launched Reserve Management Purchases (RMP)—Not-QE by another name. "
        "The plumbing has cracked.")

    # Pillar 3
    add_styled_paragraph(doc, "3. MARKET TECHNICALS: The K-Shaped Stall", 'Arial', 12, bold=True, color=OCEAN_BLUE)
    add_styled_paragraph(doc,
        "Monetary policy is now a selective tightening tool. Capital-rich megacaps are insulated "
        "by fixed-rate debt; capital-dependent SMEs are being crushed. 2026–2027 sees a $480B "
        "maturity wall in sub-IG debt.")

    add_page_break(doc)


def add_part1(doc):
    """Add Part I - The Silent Stop."""
    # Part header
    add_styled_paragraph(doc, "PART I", 'Arial', 14, bold=True, color=OCEAN_BLUE,
                        alignment=WD_ALIGN_PARAGRAPH.CENTER)
    add_styled_paragraph(doc, "THE SILENT STOP", 'Arial', 24, bold=True, color=OCEAN_BLUE,
                        alignment=WD_ALIGN_PARAGRAPH.CENTER)
    add_styled_paragraph(doc, "The U.S. Economy Has Not Slowed. It Has Stalled.",
                        'Georgia', 12, italic=True, color=NEUTRAL_GRAY,
                        alignment=WD_ALIGN_PARAGRAPH.CENTER)

    add_horizontal_line(doc)

    add_quote(doc, "GDP is the paint job. GDI is the engine. And the engine has seized.")

    add_styled_paragraph(doc,
        "The defining macro error of the current cycle is not optimism. It is misidentification of motion. "
        "Headline economic data suggests an economy that is decelerating but intact. GDP growth remains "
        "positive. Employment has softened but not collapsed. The dominant narrative is no longer "
        "'soft landing' but 'no landing.'")

    add_styled_paragraph(doc, "That narrative is wrong.", bold=True)

    add_styled_paragraph(doc,
        "The U.S. economy in early 2026 is not expanding, slowing, or contracting in the traditional "
        "cyclical sense. It is stalled—held in place by fiscal flow, accounting optics, and residual "
        "liquidity while the private-sector engine has stopped turning. This condition is what we "
        "define as The Silent Stop.")

    add_horizontal_line(doc)

    # Section I
    add_section_header(doc, "I. THE GDP–GDI SCHISM", level=2)

    add_styled_paragraph(doc,
        "If there is a single macro relationship that should anchor institutional risk assessment in 2026, "
        "it is the divergence between Gross Domestic Product (GDP) and Gross Domestic Income (GDI).")

    add_styled_paragraph(doc,
        "As of Q4 2025: Real GDP is growing at approximately +1.8% YoY. Real GDI is effectively 0.0% YoY. "
        "This is one of the largest sustained divergences on record outside of crisis periods.")

    add_chart(doc, "1.1", "Real vs Nominal GDP — Deflator masks real weakness")

    add_styled_paragraph(doc,
        "GDP captures what is spent, including government outlays financed by borrowing. GDI captures "
        "what is earned—wages, profits, interest. When GDP exceeds GDI materially and persistently, "
        "the implication is borrowed momentum.")

    add_styled_paragraph(doc,
        "In 2000, GDI rolled over before GDP. In 2007, GDI collapsed while GDP appeared stable. "
        "Each time, markets anchored to GDP were late.")

    add_page_break(doc)

    # Section II
    add_section_header(doc, "II. FISCAL DOMINANCE: WHY GDP IS LYING", level=2)

    add_styled_paragraph(doc,
        "The reason this divergence is so extreme is fiscal dominance. Government spending is growing "
        "at a pace more consistent with wartime than a mature expansion.")

    add_chart(doc, "1.2", "The Fiscal Dominance Cascade — Self-reinforcing debt spiral")

    add_styled_paragraph(doc,
        "Government spending does not create self-reinforcing economic loops. It sustains activity "
        "but does not propagate velocity. What GDP is capturing is maintenance spending, not "
        "expansionary growth.")

    add_chart(doc, "1.3", "Federal Debt Trajectory — $36T, 124% of GDP")
    add_chart(doc, "1.4", "Federal Interest Expense — 11.3% of revenue and rising")

    add_styled_paragraph(doc,
        "Private fixed investment is weak. Goods consumption is flat. Net exports are a drag. "
        "Capex outside AI-linked megacaps is contracting. This is not an expansion. It is levitation.")

    add_page_break(doc)

    # Section III
    add_section_header(doc, "III. GDI TELLS THE TRUTH", level=2)

    add_styled_paragraph(doc,
        "GDI exposes what GDP conceals. Real wage growth for the median worker has stalled. "
        "Corporate profits are bifurcated: megacaps remain strong while SMEs face margin compression.")

    add_chart(doc, "1.5", "Wealth Distribution — K-shaped reality in balance sheets")

    add_styled_paragraph(doc,
        "Tax receipts have decelerated. This is why sentiment surveys remain weak despite 'strong' "
        "headline data. Households experience the economy through income, not GDP.")

    add_page_break(doc)

    # Section IV
    add_section_header(doc, "IV. LABOR: THE ECONOMY IS FREEZING, NOT FIRING", level=2)

    add_styled_paragraph(doc,
        "The defining feature of the 2025–2026 labor market is immobility.")

    add_chart(doc, "1.6", "Employment Diffusion Index — At neutral/contraction boundary")

    add_styled_paragraph(doc,
        "The quits rate has collapsed below pre-pandemic levels. Temporary help has contracted sharply. "
        "Average weekly hours are at cycle lows. Workers are not quitting because they cannot find "
        "better jobs. Employers are not hiring aggressively.")

    add_chart(doc, "1.7", "Labor Fragility Index — Back to late-2019 warning levels")
    add_chart(doc, "1.8", "LFI Component Breakdown — Quits, hires/quits, long-duration")

    add_styled_paragraph(doc,
        "This creates the false impression of stability.")

    add_chart(doc, "1.9", "Job Cuts vs Initial Claims — Cuts at 2008 levels, claims suppressed")

    add_styled_paragraph(doc,
        "Historically, layoffs rise after this phase. The freeze precedes the break.", italic=True)

    add_page_break(doc)

    # Section V
    add_section_header(doc, "V. THE CONSUMER: FROM LIQUIDITY TO SOLVENCY", level=2)

    add_styled_paragraph(doc,
        "For three years, consumer resilience was explained by excess savings. That explanation is now obsolete.")

    add_chart(doc, "1.10", "Excess Savings Depletion — Bottom 80% exhausted June 2023")

    add_styled_paragraph(doc,
        "Excess savings are gone for the bottom 80%. Savings rates are below pre-pandemic norms. "
        "Credit card balances are at record highs. Delinquencies are accelerating.")

    add_styled_paragraph(doc, "This is no longer a liquidity story. It is a solvency story.", bold=True)

    add_chart(doc, "1.11", "Consumer Credit vs Savings — Rate at 4.0% vs 8.5% average")
    add_chart(doc, "1.12", "Two-Speed Consumer — Bifurcation in spending patterns")
    add_chart(doc, "1.13", "Subprime Auto Delinquencies — 6.67% exceeds 2008 peak")

    add_page_break(doc)

    # Section VI
    add_section_header(doc, "VI. THE K-SHAPED COST OF CAPITAL", level=2)

    add_styled_paragraph(doc,
        "Monetary policy has not tightened evenly. It has segmented the economy. Large corporations "
        "refinanced at historically low rates in 2020-2021. SMEs rely on floating-rate credit. "
        "Their effective costs are now in high single digits.")

    add_chart(doc, "1.14", "Commercial Real Estate Delinquencies — Office at 11.76% exceeds 2008")

    add_styled_paragraph(doc,
        "Markets are pricing the top decile of balance sheets and extrapolating that strength onto "
        "the entire economy. That extrapolation is incorrect.")

    add_page_break(doc)

    # Section VII & VIII
    add_section_header(doc, "VII. DEFINING THE SILENT STOP", level=2)

    add_styled_paragraph(doc, "We define The Silent Stop as an economic condition where:")
    add_bullet(doc, "Headline growth remains positive due to fiscal flow")
    add_bullet(doc, "Private income growth is flat or negative")
    add_bullet(doc, "Labor mobility collapses before unemployment rises")
    add_bullet(doc, "Consumer spending is maintained through credit, not income")
    add_bullet(doc, "Capital formation outside large incumbents stalls")

    add_styled_paragraph(doc,
        "The danger is not that markets are wrong today. The danger is that they are pricing "
        "continuity in a regime defined by fragility.")

    add_horizontal_line(doc)

    add_section_header(doc, "VIII. WHY THIS MATTERS", level=2)

    add_styled_paragraph(doc, "Stall regimes do not resolve gently. They resolve when one of three things breaks:")
    add_bullet(doc, "Funding transmission")
    add_bullet(doc, "Labor confidence")
    add_bullet(doc, "Credit availability")

    add_styled_paragraph(doc,
        "Parts II–III will show why the system is now vulnerable on all three fronts.", italic=True)

    add_page_break(doc)


def add_part2(doc):
    """Add Part II - Monetary Mechanics."""
    add_styled_paragraph(doc, "PART II", 'Arial', 14, bold=True, color=OCEAN_BLUE,
                        alignment=WD_ALIGN_PARAGRAPH.CENTER)
    add_styled_paragraph(doc, "MONETARY MECHANICS", 'Arial', 24, bold=True, color=OCEAN_BLUE,
                        alignment=WD_ALIGN_PARAGRAPH.CENTER)
    add_styled_paragraph(doc, "Reserve Scarcity & The 'Not-QE' Regime",
                        'Georgia', 12, italic=True, color=NEUTRAL_GRAY,
                        alignment=WD_ALIGN_PARAGRAPH.CENTER)

    add_horizontal_line(doc)

    add_quote(doc, "Rates tell you intent. Reserves tell you constraint.")

    add_styled_paragraph(doc,
        "Part I established the macro reality: the private economy has stalled beneath fiscal optics. "
        "Part II explains how that stall was masked, why the masking mechanism is exhausted, and why "
        "the Fed has already crossed the line from tightening to covert stabilization.")

    # Priority 1 Chart: TGA Balance
    add_priority1_chart(doc, "chart_tga_balance.png",
                       "Figure P1.1: TGA Balance History — Treasury cash volatility as liquidity lever")

    # Priority 1 Chart: Fed Balance Sheet
    add_priority1_chart(doc, "chart_fed_balance_sheet.png",
                       "Figure P1.2: Fed Balance Sheet Composition — RRP offset now exhausted")

    add_horizontal_line(doc)

    # Section I
    add_section_header(doc, "I. RESERVE SCARCITY REGIME", level=2)

    add_styled_paragraph(doc,
        "Reserve adequacy is not about absolute level. It is about distribution, behavior, "
        "and willingness to deploy.")

    add_chart(doc, "2.1", "Bank Reserves vs GDP — 12.5%, approaching Sept 2019 crisis level")

    add_styled_paragraph(doc,
        "In 2019, repo markets broke with reserves above $1T. Today the system is more constrained: "
        "Higher G-SIB surcharges, lower dealer elasticity, far greater Treasury supply.")

    # Priority 1 Chart: Fed Balance Sheet
    add_priority1_chart(doc, "chart_fed_balance_sheet.png",
                       "Figure P1.2: Fed Balance Sheet Composition — RRP offset now exhausted")

    add_page_break(doc)

    # Section II
    add_section_header(doc, "II. TREASURY SUPPLY CHALLENGE", level=2)

    add_styled_paragraph(doc,
        "With RRP gone, the Treasury General Account becomes a direct liquidity weapon. "
        "TGA up means reserves down.")

    add_chart(doc, "2.2", "Treasury Maturity Wall — $12.7T in 36 months (3x average)")
    add_chart(doc, "2.3", "Treasury Issuance by Tenor")

    add_styled_paragraph(doc,
        "Treasury cash management is now de facto monetary policy.")

    add_page_break(doc)

    # Section III
    add_section_header(doc, "III. FUNDING MARKET STRESS SIGNALS", level=2)

    add_styled_paragraph(doc,
        "In reserve-constrained regimes, small spreads matter. These are the earliest indicators "
        "of binding constraints.")

    add_chart(doc, "2.4", "SOFR-EFFR Spread — 6.1 bps in warning zone")
    add_chart(doc, "2.5", "Repo Dispersion — Widening indicates dealer stress")

    add_styled_paragraph(doc,
        "Plumbing breaks do not start with headlines. They start with basis points.", italic=True)

    add_page_break(doc)

    # Section IV
    add_section_header(doc, "IV. STANDING REPO FACILITY", level=2)

    add_chart(doc, "2.6", "Standing Repo Facility Usage — Record $48.2B October 31, 2025")
    add_chart(doc, "2.7", "SRF Usage Frequency Distribution")

    add_styled_paragraph(doc,
        "When both floor tools (RRP) and ceiling tools (SRF) are implicated in the same cycle, "
        "the system is operating at the edge.")

    add_page_break(doc)

    # Section V
    add_section_header(doc, "V. DEALER BALANCE SHEET CONSTRAINTS", level=2)

    add_chart(doc, "2.8", "Primary Dealer Balance Sheet — $94B inventory, 78% of SLR")
    add_chart(doc, "2.9", "Treasury Basis Dynamics — Below unwind threshold")

    add_horizontal_line(doc)

    add_section_header(doc, "VI. RESERVE MANAGEMENT PURCHASES (RMPs)", level=2)

    add_quote(doc, "RMPs exist because QT hit the wall.")

    add_styled_paragraph(doc,
        "RMPs are outright Treasury purchases to add reserves. Framed as technical and temporary. "
        "Mechanically: they increase reserves. Period. They are QT offsets.")

    add_styled_paragraph(doc, "If QT were harmless, RMPs would not exist.", bold=True)

    add_page_break(doc)


def add_part3(doc):
    """Add Part III - Market Technicals."""
    add_styled_paragraph(doc, "PART III", 'Arial', 14, bold=True, color=OCEAN_BLUE,
                        alignment=WD_ALIGN_PARAGRAPH.CENTER)
    add_styled_paragraph(doc, "MARKET TECHNICALS", 'Arial', 24, bold=True, color=OCEAN_BLUE,
                        alignment=WD_ALIGN_PARAGRAPH.CENTER)
    add_styled_paragraph(doc, "Liquidity-Driven Price Action & Positioning",
                        'Georgia', 12, italic=True, color=NEUTRAL_GRAY,
                        alignment=WD_ALIGN_PARAGRAPH.CENTER)

    add_horizontal_line(doc)

    add_quote(doc, "When reserves are the constraint, markets stop trending and start snapping.")

    # Section I
    add_section_header(doc, "I. TREASURY AUCTION STRESS", level=2)

    add_chart(doc, "3.1", "Auction Tails Time Series — Frequency increasing")
    add_chart(doc, "3.2", "Auction Tails Scatter — Long-duration showing most stress")
    add_chart(doc, "3.3", "Treasury Auction Deviation Analysis")

    add_page_break(doc)

    # Section II
    add_section_header(doc, "II. YIELD CURVE DYNAMICS", level=2)

    add_chart(doc, "3.4", "Treasury Yield Curve Shape Analysis")
    add_chart(doc, "3.5", "Yield Curve Heatmap — Historical configurations")
    add_chart(doc, "3.6", "Yield Curve Repricing Path — 10Y to 5.10% (+95 bps)")
    add_chart(doc, "3.7", "Ten-Year Yield Scenario Analysis")

    # Priority 1 Chart: Real Rates
    add_priority1_chart(doc, "chart_real_rates.png",
                       "Figure P1.3: Real Rates & Breakevens — Inflation expectations regime")

    add_page_break(doc)

    # Section III
    add_section_header(doc, "III. CREDIT MARKET ANALYSIS", level=2)

    add_styled_paragraph(doc,
        "Credit spreads remain tight because carry is attractive and RMPs suppress systemic stress. "
        "But credit is downstream of reserves.")

    add_chart(doc, "3.8", "Credit Spread Percentile Gauges — HY at 28th percentile")
    add_chart(doc, "3.9", "Credit Spread Waterfall — Decomposition")
    add_chart(doc, "3.10", "Credit Impulse — Leading indicator for real economy")

    add_styled_paragraph(doc,
        "Credit should be treated as late-cycle convex risk, not yield pickup.", italic=True)

    add_page_break(doc)

    # Section IV
    add_section_header(doc, "IV. CROSS-ASSET SIGNALS", level=2)

    add_chart(doc, "3.11", "Cross-Asset Correlations — Stock-bond at -0.48")

    add_styled_paragraph(doc,
        "One dangerous assumption: that duration reliably hedges equity risk. When term premium rises "
        "and liquidity support is targeted—equities and bonds can sell off together.")

    add_chart(doc, "3.12", "Foreign Official Treasury Holdings — China declining structurally")

    # Priority 1 Charts
    add_priority1_chart(doc, "chart_dollar_index.png",
                       "Figure P1.4: Trade-Weighted Dollar Index — USD funding demand")
    add_priority1_chart(doc, "chart_move_vs_vix.png",
                       "Figure P1.5: MOVE vs VIX — Rates volatility leads equity volatility")

    # Priority 1 Chart: MMF Flows
    add_priority1_chart(doc, "chart_mmf_flows.png",
                       "Figure P1.6: Money Market Fund Flows — Where RRP cash migrated")

    # Priority 1 Chart: SLOOS
    add_priority1_chart(doc, "chart_sloos.png",
                       "Figure P1.7: Bank Lending Standards (SLOOS) — Credit tightening evidence")

    add_horizontal_line(doc)

    # Synthesis
    add_section_header(doc, "V. SYNTHESIS: HOW RMP REGIMES BREAK", level=2)

    add_styled_paragraph(doc, "The sequence:")
    for i, item in enumerate([
        "RRP is gone", "Reserves are scarce", "QT is capped by RMPs",
        "Markets extrapolate stability", "Positioning accumulates",
        "A funding or issuance shock hits", "RMPs respond with lag",
        "Liquidity evaporates", "Markets gap"
    ], 1):
        add_bullet(doc, f"{i}. {item}")

    add_styled_paragraph(doc,
        "Stabilization in reserve-scarce regimes is discrete, reactive, and uneven.", bold=True)

    add_page_break(doc)


def add_conclusion(doc):
    """Add Conclusion."""
    add_styled_paragraph(doc, "CONCLUSION", 'Arial', 14, bold=True, color=OCEAN_BLUE,
                        alignment=WD_ALIGN_PARAGRAPH.CENTER)
    add_styled_paragraph(doc, "THE SILENT STOP, REVISITED", 'Arial', 24, bold=True, color=OCEAN_BLUE,
                        alignment=WD_ALIGN_PARAGRAPH.CENTER)
    add_styled_paragraph(doc, "Why This Cycle Ends with Discontinuity, Not Resolution",
                        'Georgia', 12, italic=True, color=NEUTRAL_GRAY,
                        alignment=WD_ALIGN_PARAGRAPH.CENTER)

    add_horizontal_line(doc)

    add_chart(doc, "4.1", "The 16-Week Stress Window — Critical event calendar")

    add_styled_paragraph(doc,
        "The defining error of the current macro moment is not optimism. It is misdiagnosis. "
        "The economy has not normalized. It has stalled.")

    add_styled_paragraph(doc,
        "GDI is flat. Labor mobility has collapsed. Hours are at cycle lows. The consumer has shifted "
        "to solvency stress. SMEs face prohibitive capital costs. These are not late-cycle warnings. "
        "They are late-cycle facts.")

    add_horizontal_line(doc)

    # Core synthesis points
    add_section_header(doc, "THE CORE SYNTHESIS", level=2)

    for num, (title, text) in enumerate([
        ("LOSS OF VELOCITY", "In a normal slowdown, the economy moves forward more slowly. "
         "In this regime, motion itself has broken down."),
        ("FISCAL MASKING", "Government spending has prevented contraction. It has not restored "
         "private-sector momentum. When fiscal impulse fades, the private economy will be exposed."),
        ("LIQUIDITY CONSTRAINT", "The liquidity buffer has been exhausted. RMPs are an admission "
         "that QT reached its limit. Liquidity support is now conditional, reactive, and uneven."),
        ("PRICING CONTINUITY IN FRAGILITY", "Asset prices assume the current equilibrium can persist. "
         "History suggests such regimes end when one transmission channel fails.")
    ], 1):
        add_styled_paragraph(doc, f"{num}. {title}", 'Arial', 12, bold=True, color=OCEAN_BLUE)
        add_styled_paragraph(doc, text)

    add_horizontal_line(doc)

    # Investor implications
    add_section_header(doc, "WHAT THIS MEANS FOR INVESTORS", level=2)

    add_styled_paragraph(doc,
        "The dominant risk is not being early on a slowdown. It is being structurally exposed to "
        "nonlinear repricing.")

    add_bullet(doc, "Carry trades work until they don't")
    add_bullet(doc, "Diversification fails intermittently")
    add_bullet(doc, "Liquidity disappears when most needed")
    add_bullet(doc, "The cost of being wrong is asymmetric")

    add_styled_paragraph(doc,
        "Portfolios should be built to survive gaps, not optimize for smooth paths. "
        "The goal is not to avoid risk. It is to avoid being forced.", bold=True)

    add_horizontal_line(doc)

    # Final thought
    add_section_header(doc, "FINAL THOUGHT", level=2)

    add_styled_paragraph(doc,
        "This is not a soft landing. It is not a hard landing. It is not a reacceleration. "
        "It is levitation over a vacuum.")

    add_styled_paragraph(doc,
        "The economy has stopped moving forward, but asset prices have not yet reconciled with that reality. "
        "That gap will close. When it does, the adjustment will not be linear.")

    add_styled_paragraph(doc,
        "The Silent Stop does not announce itself. It reveals itself—all at once.", bold=True)

    add_styled_paragraph(doc, "Position accordingly.")

    add_horizontal_line(doc)

    # Sign-off
    for _ in range(2):
        doc.add_paragraph()

    add_styled_paragraph(doc,
        "That's our view from the Watch. Until next time, we'll be sure to keep the light on....",
        'Georgia', 11, italic=True, color=NEUTRAL_GRAY, alignment=WD_ALIGN_PARAGRAPH.CENTER)

    for _ in range(2):
        doc.add_paragraph()

    add_styled_paragraph(doc, "— Bob Sheehan, CFA, CMT", 'Arial', 12, bold=True,
                        alignment=WD_ALIGN_PARAGRAPH.CENTER)
    add_styled_paragraph(doc, "Founder & Chief Investment Officer", 'Arial', 11,
                        alignment=WD_ALIGN_PARAGRAPH.CENTER)
    add_styled_paragraph(doc, "Lighthouse Macro", 'Arial', 11, color=OCEAN_BLUE,
                        alignment=WD_ALIGN_PARAGRAPH.CENTER)

    for _ in range(2):
        doc.add_paragraph()

    add_styled_paragraph(doc, "MACRO, ILLUMINATED.", 'Arial', 14, bold=True, color=OCEAN_BLUE,
                        alignment=WD_ALIGN_PARAGRAPH.CENTER)


def generate_report():
    """Generate the complete report."""
    print("=" * 60)
    print("GENERATING THE HORIZON REPORT - JANUARY 2026")
    print("FINAL INSTITUTIONAL VERSION")
    print("=" * 60)

    doc = create_document()

    print("\n[1/6] Adding cover page...")
    add_cover_page(doc)

    print("[2/6] Adding executive summary...")
    add_executive_summary(doc)

    print("[3/6] Adding Part I - The Silent Stop...")
    add_part1(doc)

    print("[4/6] Adding Part II - Monetary Mechanics...")
    add_part2(doc)

    print("[5/6] Adding Part III - Market Technicals...")
    add_part3(doc)

    print("[6/6] Adding Conclusion...")
    add_conclusion(doc)

    # Save
    timestamp = datetime.now().strftime("%Y%m%d_%H%M")
    output_path = OUTPUT_DIR / f"The_Horizon_Report_Jan2026_FINAL_{timestamp}.docx"

    print(f"\nSaving: {output_path}")
    doc.save(str(output_path))

    # Save as latest
    latest_path = OUTPUT_DIR / "The_Horizon_Report_Jan2026_FINAL.docx"
    doc.save(str(latest_path))

    print("\n" + "=" * 60)
    print("REPORT GENERATION COMPLETE")
    print("=" * 60)
    print(f"\nTimestamped: {output_path}")
    print(f"Latest: {latest_path}")

    # Summary
    print("\nCHART SUMMARY:")
    found = 0
    missing = 0
    for fig_num, filename in CHART_MAP.items():
        path = CHART_DIR / filename
        if path.exists():
            found += 1
        else:
            missing += 1
            print(f"  Missing: {fig_num} - {filename}")

    print(f"\n  Found: {found}/{len(CHART_MAP)} Horizon charts")

    # Check Priority 1 charts
    p1_charts = ["chart_tga_balance.png", "chart_fed_balance_sheet.png",
                 "chart_real_rates.png", "chart_dollar_index.png", "chart_move_vs_vix.png",
                 "chart_mmf_flows.png", "chart_sloos.png"]
    p1_found = sum(1 for c in p1_charts if (PRIORITY1_DIR / c).exists())
    print(f"  Found: {p1_found}/{len(p1_charts)} Priority 1 charts")

    return output_path


if __name__ == "__main__":
    generate_report()
