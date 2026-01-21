#!/usr/bin/env python3
"""
LIGHTHOUSE MACRO - WORD CHARTBOOK GENERATOR
Generates .docx instead of PDF for easy editing

Author: Bob Sheehan, CFA, CMT
Date: November 2025
"""

from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
import pandas as pd
from datetime import datetime
import os

# Paths
PROP_CHARTS = "/Users/bob/lighthouse_paywall_deck/charts/proprietary"
MACRO_CHARTS = "/Users/bob/macromicro_charts"
TV_CHARTS = "/Users/bob/tradingview_charts"
INDICATORS_CSV = "/Users/bob/lighthouse_paywall_deck/proprietary_indicators.csv"
FSI_CSV = "/Users/bob/fsi.csv"
OUTPUT_DOCX = "/Users/bob/Lighthouse_Macro_Premium_Chartbook.docx"

def load_latest_indicators():
    """Load the most recent indicator values"""
    df = pd.read_csv(INDICATORS_CSV, index_col=0, parse_dates=[0])
    latest = df.iloc[-1]

    # Load FSI data
    fsi_df = pd.read_csv(FSI_CSV, parse_dates=['Date'], index_col='Date')
    latest_fsi = fsi_df.iloc[-1]

    return {
        'date': df.index[-1].strftime('%B %d, %Y'),
        'MRI': latest['MRI'],
        'LCI': latest['LCI'],
        'LFI': latest['LFI'],
        'LDI': latest['LDI'],
        'CLG': latest['CLG'],
        'YFS': latest['YFS'],
        'SVI': latest['SVI'],
        'EMD': latest['EMD'],
        'FSI': latest_fsi['OFR FSI'],
    }

print("=" * 70)
print("LIGHTHOUSE MACRO - WORD CHARTBOOK GENERATOR")
print("=" * 70)

# Load data
print("\n[1/5] Loading indicator data...")
indicators = load_latest_indicators()
print(f"   MRI: {indicators['MRI']:.2f}σ")
print(f"   LCI: {indicators['LCI']:.2f}σ")

# Create document
print("\n[2/5] Creating Word document...")
doc = Document()

# Set default font
style = doc.styles['Normal']
font = style.font
font.name = 'Arial'
font.size = Pt(11)

# Title page
print("\n[3/5] Building cover page...")
title = doc.add_heading('LIGHTHOUSE MACRO', level=0)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER

subtitle = doc.add_heading('Premium Institutional Chartbook', level=1)
subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER

doc.add_paragraph(f"Data as of {indicators['date']}")
doc.add_paragraph()

# Intro text (simplified - no duplicate MRI)
intro = f"""Our Macro Risk Index (MRI) currently reads {indicators['MRI']:.2f}σ (near neutral).

Key factor: Liquidity Cushion Index at {indicators['LCI']:.2f}σ — critically depleted banking system liquidity.

Funding stress building (+{indicators['YFS']:.2f}σ) while labor fragility elevated (+{indicators['LFI']:.2f}σ).

This chartbook combines proprietary quantitative indicators, global macro intelligence, and technical analysis."""

doc.add_paragraph(intro)
doc.add_page_break()

# SECTION I: PROPRIETARY INDICATORS
print("\n[4/5] Adding proprietary charts...")

doc.add_heading('SECTION I: PROPRIETARY INDICATORS', level=1)

priority_charts = [
    ('MRI_Macro_Risk_Index.png', 'Macro Risk Index (MRI)'),
    ('01_LCI_Liquidity_Cushion_Index.png', 'Liquidity Cushion Index (LCI)'),
    ('02_LFI_Labor_Fragility_Index.png', 'Labor Fragility Index (LFI)'),
    ('03_LDI_Labor_Dynamism_Index.png', 'Labor Dynamism Index (LDI)'),
    ('04_CLG_Credit_Labor_Gap.png', 'Credit-Labor Gap (CLG)'),
    ('05_YFS_Yield_Funding_Stress.png', 'Yield-Funding Stress (YFS)'),
    ('06_SVI_Spread_Volatility_Imbalance.png', 'Spread-Volatility Imbalance (SVI)'),
    ('09_Payrolls_vs_Quits_Divergence.png', 'Payrolls vs Quits Divergence'),
    ('10_Hours_vs_Employment_Divergence.png', 'Hours vs Employment Divergence'),
]

for chart_file, chart_name in priority_charts:
    chart_path = os.path.join(PROP_CHARTS, chart_file)
    if os.path.exists(chart_path):
        doc.add_heading(chart_name, level=2)
        doc.add_picture(chart_path, width=Inches(6.5))
        doc.add_page_break()
        print(f"   ✓ {chart_name}")

# SECTION II: MACROMICRO
print("\n[5/5] Adding MacroMicro charts...")

doc.add_heading('SECTION II: GLOBAL MACRO INTELLIGENCE', level=1)

macro_files = sorted([f for f in os.listdir(MACRO_CHARTS) if f.endswith('.png')])

# Deduplicate
seen_charts = set()
unique_files = []
for f in macro_files:
    base_name = f.replace(' (1)', '').replace(' (2)', '').replace(' (3)', '')
    if base_name not in seen_charts:
        seen_charts.add(base_name)
        unique_files.append(f)

for chart_file in unique_files:
    chart_path = os.path.join(MACRO_CHARTS, chart_file)
    chart_name = chart_file.replace('mm-chart-', '').replace('.png', '').split('_', 1)[-1]
    doc.add_heading(chart_name.replace('-', ' '), level=2)
    doc.add_picture(chart_path, width=Inches(6.5))
    doc.add_page_break()
    print(f"   ✓ {chart_name}")

# SECTION III: TRADINGVIEW
if os.path.exists(TV_CHARTS):
    doc.add_heading('SECTION III: TECHNICAL ANALYSIS', level=1)

    tv_files = sorted([f for f in os.listdir(TV_CHARTS) if f.endswith('.png')])
    for chart_file in tv_files:
        chart_path = os.path.join(TV_CHARTS, chart_file)
        ticker = chart_file.split('_')[0]
        doc.add_heading(ticker, level=2)
        doc.add_picture(chart_path, width=Inches(6.5))
        doc.add_page_break()
        print(f"   ✓ {ticker}")

# Save
doc.save(OUTPUT_DOCX)

file_size = os.path.getsize(OUTPUT_DOCX) / (1024 * 1024)

print("\n" + "=" * 70)
print("CHARTBOOK COMPLETE!")
print("=" * 70)
print(f"\nOutput: {OUTPUT_DOCX}")
print(f"Size: {file_size:.1f} MB")
print(f"Format: Microsoft Word (.docx)")
print(f"\nCurrent MRI: {indicators['MRI']:.2f}σ")
print("\n✓ Ready for editing in Word!")
print("=" * 70)
